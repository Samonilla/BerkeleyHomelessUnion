import io
import re
import textwrap


def _case_name(case: dict) -> str:
    plaintiff_name = str((case.get("plaintiff") or {}).get("name") or "Plaintiff").strip()
    defendant = case.get("defendant") or {}
    defendant_name = str(defendant.get("name") or "Defendant").strip()
    if case.get("additional_defendants"):
        defendant_name = f"{defendant_name}, et al."
    return f"{plaintiff_name} v. {defendant_name}"


def _normalize_plain_language(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if not text:
        return ""
    text = text[0].upper() + text[1:]
    if text[-1] not in ".!?":
        text += "."
    return text


def build_guided_declaration(text: str, answers: dict) -> str:
    intro = [
        "I am the plaintiff in this action.",
        "I am submitting this declaration under penalty of perjury and state that the following is true and correct.",
    ]
    facts = []

    if answers.get("incident_date"):
        facts.append(f"On {answers['incident_date']}, the events described below occurred.")

    for key, value in answers.items():
        if key in {"incident_date", "items", "claim_amount"}:
            continue
        cleaned = _normalize_plain_language(value)
        if cleaned:
            facts.append(cleaned)

    items = answers.get("items") or []
    if items:
        item_lines = []
        for item in items:
            description = str(item.get("description") or "").strip()
            value = str(item.get("value") or "").strip()
            condition = str(item.get("condition") or "").strip() or "Unknown"
            if description and value:
                item_lines.append(
                    f"I lost {description}, which was valued at ${value}, and it was in {condition.lower()} condition when it was destroyed."
                )
            elif description:
                item_lines.append(
                    f"I lost {description}, and it was in {condition.lower()} condition when it was destroyed."
                )
        if item_lines:
            facts.append(" ".join(item_lines))

    if text.strip():
        facts.append(_normalize_plain_language(text))

    if answers.get("claim_amount"):
        facts.append(f"The total value of the property I lost is approximately ${answers['claim_amount']}.")

    paragraphs = intro + [f"{i}. {fact}" for i, fact in enumerate(facts, start=1)]
    paragraphs.append("I declare under penalty of perjury that the foregoing is true and correct.")
    return "\n\n".join(paragraphs)


def build_declaration_docx(text: str) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Declaration", level=1)
    for paragraph_text in text.split("\n\n"):
        if paragraph_text.strip():
            document.add_paragraph(paragraph_text)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_govt_claim_docx(data: dict) -> bytes:
    """Build a Gov. Code sections 905/910 claim-for-damages form as a Word document."""
    from docx import Document

    entity = (data.get("entity") or "").strip()
    document = Document()
    document.add_heading("CLAIM FOR DAMAGES AGAINST A PUBLIC ENTITY", level=1)
    to_text = f"To: Office of the City Clerk{', ' + entity if entity else ''}"
    if data.get("clerk_address"):
        to_text += f"\n{data['clerk_address']}"
    document.add_paragraph(to_text)
    document.add_paragraph(
        "This claim is presented pursuant to California Government Code "
        "sections 905, 910, and 910.2."
    )

    def _row(label, value):
        p = document.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value if value else "____________________________________________")

    _row("1. Claimant name", data.get("claimant_name"))
    _row("2. Mailing address (send all notices here)", data.get("claimant_address"))
    _row("3. Phone", data.get("claimant_phone"))
    _row("4. Date of occurrence", data.get("incident_date"))
    _row("5. Place of occurrence", data.get("incident_location"))
    _row("6. Circumstances of the occurrence", data.get("description"))
    _row("7. General description of the injury, damage, or loss", data.get("description"))
    _row("8. Names of public employees or agencies causing the loss, if known", data.get("employees"))

    items = data.get("items") or []
    if items:
        p = document.add_paragraph()
        p.add_run("Itemized property destroyed or taken: ").bold = True
        for item in items:
            desc = str(item.get("description") or "").strip()
            val = str(item.get("value") or "").replace("$", "").strip()
            cond = str(item.get("condition") or "").strip()
            line = f"- {desc}"
            if cond:
                line += f" (condition: {cond})"
            if val:
                line += f" - ${val}"
            document.add_paragraph(line)

    amount_raw = (data.get("amount") or "").replace("$", "").replace(",", "").strip()
    try:
        amount_val = float(amount_raw)
    except ValueError:
        amount_val = None
    if amount_val is not None and amount_val > 10000:
        _row(
            "9. Amount claimed",
            "The amount claimed exceeds $10,000 and this would be a limited "
            "civil case. (Gov. Code section 910(f).)",
        )
    else:
        _row(
            "9. Amount claimed as of presentation, with basis of computation",
            (f"${amount_raw} - the value of the claimant's personal property "
             "destroyed or taken, as described above.") if amount_raw else "",
        )

    document.add_paragraph("")
    document.add_paragraph(
        "I declare under penalty of perjury under the laws of the State of "
        "California that the foregoing is true and correct."
    )
    document.add_paragraph("Dated: ____________________")
    document.add_paragraph(
        f"Signature: ____________________    {data.get('claimant_name') or ''}"
    )

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_subpoena_attachments_docx(data: dict) -> bytes:
    """Build a Word document containing only attachment pages for a subpoena."""
    from docx import Document

    document = Document()
    document.add_heading("Attachment to Small Claims Subpoena", level=1)

    case_caption = (data.get("case_caption") or "").strip()
    if case_caption:
        document.add_paragraph(f"Case caption: {case_caption}")

    document.add_paragraph(
        "Attach these pages behind the subpoena as the list of documents and "
        "records requested."
    )

    recipient = (data.get("to") or "").strip()
    custodian = (data.get("custodian") or "").strip()
    service_location = (data.get("service_location") or "").strip()

    if recipient or custodian or service_location:
        document.add_heading("Records Requested From", level=2)
        if recipient:
            document.add_paragraph(f"Person or agency: {recipient}")
        if custodian:
            document.add_paragraph(f"Custodian of records: {custodian}")
        if service_location:
            document.add_paragraph(f"Service address: {service_location}")

    requests = [str(request).strip() for request in (data.get("requests") or []) if str(request).strip()]
    document.add_heading("Requested Documents and Things", level=2)
    if requests:
        for index, request in enumerate(requests, start=1):
            document.add_paragraph(f"{index}. {request}")
    else:
        document.add_paragraph("No specific requests were listed.")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_exhibit_covers_pdf(exhibits: list[dict], case: dict) -> bytes:
    """Build exhibit face pages: one page per exhibit with label + description."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    case_name = _case_name(case)

    for idx, ex in enumerate(exhibits):
        label = f"Exhibit {chr(65 + idx)}"
        description = str(ex.get("description") or "").strip() or "(No description provided)"

        c.setFont("Helvetica-Bold", 24)
        c.drawCentredString(width / 2, height - 120, label.upper())

        c.setFont("Helvetica", 12)
        c.drawString(72, height - 165, f"Case: {case_name}")
        c.drawString(72, height - 185, f"Description: {description[:180]}")

        c.setFont("Helvetica", 11)
        y = height - 230
        for line in textwrap.wrap(description, width=95):
            c.drawString(72, y, line)
            y -= 16
            if y < 100:
                break

        c.setFont("Helvetica-Oblique", 10)
        c.drawString(72, 72, "This is an exhibit face page automatically generated by the intake app.")
        c.showPage()

    c.save()
    return buf.getvalue()
