import io
import re
from datetime import datetime

import streamlit as st
from docx import Document


def _fallback_slug(value: str) -> str:
    token = re.sub(r"[^a-z0-9]+", "_", str(value or "").strip().lower()).strip("_")
    return token or "ag_complaint"


def build_ag_complaint_docx(data: dict) -> bytes:
    """Build a California AG complaint draft letter as a Word document."""
    document = Document()
    document.add_heading("California Attorney General Complaint Draft", level=1)
    document.add_paragraph("To: California Department of Justice, Public Inquiry Unit")
    document.add_paragraph(f"Date: {datetime.now().strftime('%m/%d/%Y')}")

    def _row(label, value):
        p = document.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run((value or "").strip() or "____________________________________________")

    _row("Complainant name", data.get("name", ""))
    _row("Phone", data.get("phone", ""))
    _row("Email", data.get("email", ""))
    _row("Address", data.get("address", ""))
    _row("Entity or facility complained about", data.get("entity", ""))
    _row("Location", data.get("location", ""))
    _row("Date(s) of incident", data.get("incident_date", ""))
    _row("Subject", data.get("subject", ""))

    document.add_paragraph("")
    p = document.add_paragraph()
    p.add_run("Complaint details:").bold = True
    document.add_paragraph((data.get("details") or "").strip() or "(Add details here.)")

    p = document.add_paragraph()
    p.add_run("Requested action:").bold = True
    document.add_paragraph((data.get("requested_action") or "").strip() or "(Add requested action here.)")

    document.add_paragraph("")
    document.add_paragraph(
        "I declare under penalty of perjury under the laws of the State of "
        "California that the information provided above is true and correct "
        "to the best of my knowledge."
    )
    document.add_paragraph("Signature: ______________________________")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def render_ag_complaint_form(
    prefix: str,
    *,
    default_entity: str = "",
    default_location: str = "",
    default_subject: str = "",
    slug_fn=None,
) -> None:
    slug_fn = slug_fn or _fallback_slug

    c1, c2 = st.columns(2)
    with c1:
        complainant_name = st.text_input("Full Name *", key=f"{prefix}_name")
        complainant_phone = st.text_input("Phone", key=f"{prefix}_phone")
        complainant_email = st.text_input("Email", key=f"{prefix}_email")
    with c2:
        complainant_address = st.text_input("Mailing Address", key=f"{prefix}_address")
        entity = st.text_input(
            "Entity or Facility Complained About *",
            value=default_entity,
            key=f"{prefix}_entity",
        )
        location = st.text_input(
            "Location",
            value=default_location,
            key=f"{prefix}_location",
        )

    incident_date = st.text_input(
        "Date(s) of incident",
        placeholder="MM/DD/YYYY or date range",
        key=f"{prefix}_incident_date",
    )
    subject = st.text_input(
        "Complaint Subject *",
        value=default_subject,
        key=f"{prefix}_subject",
    )
    details = st.text_area(
        "What happened? *",
        key=f"{prefix}_details",
        height=170,
        placeholder="Describe what happened, who was involved, and any harm caused.",
    )
    requested_action = st.text_area(
        "What action do you want the AG to take?",
        key=f"{prefix}_requested_action",
        height=110,
        placeholder="Example: investigate the shelter, require corrective action, and ensure resident safety.",
    )

    if st.button("Generate AG Complaint Draft", type="primary", use_container_width=True, key=f"{prefix}_gen_ag"):
        if not complainant_name.strip() or not subject.strip() or not details.strip() or not entity.strip():
            st.warning("Please fill in Name, Entity, Subject, and What happened before generating.")
        else:
            payload = {
                "name": complainant_name.strip(),
                "phone": complainant_phone.strip(),
                "email": complainant_email.strip(),
                "address": complainant_address.strip(),
                "entity": entity.strip(),
                "location": location.strip(),
                "incident_date": incident_date.strip(),
                "subject": subject.strip(),
                "details": details.strip(),
                "requested_action": requested_action.strip(),
            }
            st.session_state[f"{prefix}_ag_docx"] = build_ag_complaint_docx(payload)
            st.session_state[f"{prefix}_ag_txt"] = (
                f"Complainant: {payload['name']}\n"
                f"Phone: {payload['phone']}\n"
                f"Email: {payload['email']}\n"
                f"Address: {payload['address']}\n"
                f"Entity: {payload['entity']}\n"
                f"Location: {payload['location']}\n"
                f"Date(s): {payload['incident_date']}\n"
                f"Subject: {payload['subject']}\n\n"
                f"Complaint details:\n{payload['details']}\n\n"
                f"Requested action:\n{payload['requested_action']}\n"
            ).strip()
            st.session_state[f"{prefix}_ag_filebase"] = slug_fn(payload["name"] or "ag_complaint")

    if st.session_state.get(f"{prefix}_ag_docx"):
        st.success("Complaint draft generated. Review, edit if needed, then submit to the California AG.")
        filebase = st.session_state.get(f"{prefix}_ag_filebase", "ag_complaint")
        st.download_button(
            "⬇️ Download AG Complaint (Word)",
            data=st.session_state[f"{prefix}_ag_docx"],
            file_name=f"{filebase}_ag_complaint.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"{prefix}_dl_ag_docx",
        )
        st.download_button(
            "⬇️ Download AG Complaint (Text)",
            data=st.session_state[f"{prefix}_ag_txt"].encode(),
            file_name=f"{filebase}_ag_complaint.txt",
            mime="text/plain",
            use_container_width=True,
            key=f"{prefix}_dl_ag_txt",
        )


def render_ag_complaints_ui(initial_tab: str = "general", slug_fn=None) -> None:
    st.subheader("California AG Complaint Autofiler")
    st.caption(
        "Generate a ready-to-review complaint draft for submission to the "
        "California Attorney General."
    )

    if initial_tab == "sam_jones":
        ag_tab_sam_jones, ag_tab_general = st.tabs([
            "Sam Jones Complaints",
            "General AG Complaint",
        ])
    else:
        ag_tab_general, ag_tab_sam_jones = st.tabs([
            "General AG Complaint",
            "Sam Jones Complaints",
        ])

    with ag_tab_general:
        st.markdown("**General California AG complaint**")
        render_ag_complaint_form(
            "ag_general",
            default_subject="Complaint regarding shelter or public service conditions",
            slug_fn=slug_fn,
        )

    with ag_tab_sam_jones:
        st.markdown("**Sam Jones Shelter complaint**")
        st.caption(
            "This tab is prefilled for Sam Jones Hall so residents and "
            "advocates can quickly draft AG complaints specific to that shelter."
        )
        render_ag_complaint_form(
            "ag_sam_jones",
            default_entity="Sam Jones Hall Homeless Shelter",
            default_location="Santa Rosa, California",
            default_subject="Complaint regarding conditions and treatment at Sam Jones Hall Homeless Shelter",
            slug_fn=slug_fn,
        )
