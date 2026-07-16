"""
Oakland Encampment Small Claims — Streamlit UI

Run:  streamlit run app.py
"""

import contextlib
import io
import json
import os
import re
import sys
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path

import pandas as pd
import streamlit as st
from dateutil import parser as _dateutil
from docx import Document

HERE = Path(__file__).parent
sys.path.insert(0, str(HERE))

from fill_forms import (
    fill_sc100, fill_fw001, fill_fw003, fill_sc112a, fill_sc150,
    fill_sc105, fill_sc107, fill_sc109, fill_sc100a_for_party,
    validate_case, has_postponement, DEFENDANT_DEFAULTS,
    _SC107_DEFAULT_GOOD_CAUSE, _SC107_DEFAULT_MATERIALITY,
)
from courts import ALL_COUNTIES, courthouses_for_county, court_info_string
from defendants import ALL_CITIES, defendant_info
from storage import (
    case_dirs as _case_dirs,
    primary_cases_dir as _primary_cases_dir,
    slug as _slug,
    capture_case_record as _capture_case_record,
    load_cases as _load_case_files,
)

_META_SC100 = str(HERE / "field_meta" / "sc100_fields.json")
_META_FW001 = str(HERE / "field_meta" / "fw001_fields.json")
_TPL = HERE / "templates"


def _normalize_plain_language(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if not text:
        return ""
    text = text[0].upper() + text[1:]
    if text[-1] not in ".!?":
        text += "."
    return text


def _build_guided_declaration(text: str, answers: dict) -> str:
    intro = [
        "I am the plaintiff in this action.",
        "I am submitting this declaration under penalty of perjury and state that the following is true and correct.",
    ]
    facts = []

    if answers.get("incident_date"):
        facts.append(f"On {answers['incident_date']}, the events described below occurred.")

    for key, value in answers.items():
        if key in {"incident_date", "items", "claim_amount"}:
            continue
        cleaned = _normalize_plain_language(value)
        if cleaned:
            facts.append(cleaned)

    items = answers.get("items") or []
    if items:
        item_lines = []
        for item in items:
            description = str(item.get("description") or "").strip()
            value = str(item.get("value") or "").strip()
            condition = str(item.get("condition") or "").strip() or "Unknown"
            if description and value:
                item_lines.append(
                    f"I lost {description}, which was valued at ${value}, and it was in {condition.lower()} condition when it was destroyed."
                )
            elif description:
                item_lines.append(
                    f"I lost {description}, and it was in {condition.lower()} condition when it was destroyed."
                )
        if item_lines:
            facts.append(" ".join(item_lines))

    if text.strip():
        facts.append(_normalize_plain_language(text))

    if answers.get("claim_amount"):
        facts.append(f"The total value of the property I lost is approximately ${answers['claim_amount']}.")

    paragraphs = intro + [f"{i}. {fact}" for i, fact in enumerate(facts, start=1)]
    paragraphs.append("I declare under penalty of perjury that the foregoing is true and correct.")
    return "\n\n".join(paragraphs)


def _build_declaration_docx(text: str) -> bytes:
    document = Document()
    document.add_heading("Declaration", level=1)
    for paragraph_text in text.split("\n\n"):
        if paragraph_text.strip():
            document.add_paragraph(paragraph_text)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _build_govt_claim_docx(data: dict) -> bytes:
    """Build a Gov. Code §§ 905/910 claim-for-damages form as a Word document."""
    entity = (data.get("entity") or "").strip()
    document = Document()
    document.add_heading("CLAIM FOR DAMAGES AGAINST A PUBLIC ENTITY", level=1)
    to_text = f"To: Office of the City Clerk{', ' + entity if entity else ''}"
    if data.get("clerk_address"):
        to_text += f"\n{data['clerk_address']}"
    document.add_paragraph(to_text)
    document.add_paragraph(
        "This claim is presented pursuant to California Government Code "
        "sections 905, 910, and 910.2."
    )

    def _row(label, value):
        p = document.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value if value else "____________________________________________")

    _row("1. Claimant name", data.get("claimant_name"))
    _row("2. Mailing address (send all notices here)", data.get("claimant_address"))
    _row("3. Phone", data.get("claimant_phone"))
    _row("4. Date of occurrence", data.get("incident_date"))
    _row("5. Place of occurrence", data.get("incident_location"))
    _row("6. Circumstances of the occurrence", data.get("description"))
    _row("7. General description of the injury, damage, or loss", data.get("description"))
    _row("8. Names of public employees or agencies causing the loss, if known",
         data.get("employees"))

    items = data.get("items") or []
    if items:
        p = document.add_paragraph()
        p.add_run("Itemized property destroyed or taken: ").bold = True
        for item in items:
            desc = str(item.get("description") or "").strip()
            val = str(item.get("value") or "").replace("$", "").strip()
            cond = str(item.get("condition") or "").strip()
            line = f"– {desc}"
            if cond:
                line += f" (condition: {cond})"
            if val:
                line += f" — ${val}"
            document.add_paragraph(line)

    amount_raw = (data.get("amount") or "").replace("$", "").replace(",", "").strip()
    try:
        amount_val = float(amount_raw)
    except ValueError:
        amount_val = None
    if amount_val is not None and amount_val > 10000:
        _row(
            "9. Amount claimed",
            "The amount claimed exceeds $10,000 and this would be a limited "
            "civil case. (Gov. Code § 910(f).)",
        )
    else:
        _row(
            "9. Amount claimed as of presentation, with basis of computation",
            (f"${amount_raw} — the value of the claimant's personal property "
             "destroyed or taken, as described above.") if amount_raw else "",
        )

    document.add_paragraph("")
    document.add_paragraph(
        "I declare under penalty of perjury under the laws of the State of "
        "California that the foregoing is true and correct."
    )
    document.add_paragraph("Dated: ____________________")
    document.add_paragraph(
        f"Signature: ____________________    {data.get('claimant_name') or ''}"
    )

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _build_subpoena_attachments_docx(data: dict) -> bytes:
    """Build a Word document containing only attachment pages for a subpoena."""
    document = Document()
    document.add_heading("Attachment to Small Claims Subpoena", level=1)

    case_caption = (data.get("case_caption") or "").strip()
    if case_caption:
        document.add_paragraph(f"Case caption: {case_caption}")

    document.add_paragraph(
        "Attach these pages behind the subpoena as the list of documents and "
        "records requested."
    )

    recipient = (data.get("to") or "").strip()
    custodian = (data.get("custodian") or "").strip()
    service_location = (data.get("service_location") or "").strip()

    if recipient or custodian or service_location:
        document.add_heading("Records Requested From", level=2)
        if recipient:
            document.add_paragraph(f"Person or agency: {recipient}")
        if custodian:
            document.add_paragraph(f"Custodian of records: {custodian}")
        if service_location:
            document.add_paragraph(f"Service address: {service_location}")

    requests = [
        str(request).strip() for request in (data.get("requests") or []) if str(request).strip()
    ]
    document.add_heading("Requested Documents and Things", level=2)
    if requests:
        for index, request in enumerate(requests, start=1):
            document.add_paragraph(f"{index}. {request}")
    else:
        document.add_paragraph("No specific requests were listed.")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ─── Local jurisdiction claim form (uploaded PDF) auto-fill ──────────────────

# Order matters: more specific patterns first (e.g. email before address,
# incident date before generic date). Matched against field name + tooltip.
_LOCAL_CLAIM_PATTERNS = [
    ("claimant_email",    re.compile(r"e-?mail", re.I)),
    ("claimant_phone",    re.compile(r"phone|telephone", re.I)),
    ("claimant_name",     re.compile(r"(claimant|your|full|print|last|first).{0,20}name|name.{0,10}of.{0,10}claimant|^name\b", re.I)),
    ("claimant_address",  re.compile(r"address", re.I)),
    ("incident_date",     re.compile(
        r"date.{0,30}(incident|occurr|loss|injur|accident|damage|event)"
        r"|(incident|occurr|loss|injur|accident|damage|event).{0,30}date"
        r"|when did", re.I)),
    ("incident_location", re.compile(r"location|place|where", re.I)),
    ("amount",            re.compile(r"amount|total.{0,15}claim|damages?\b", re.I)),
    ("description",       re.compile(
        r"describe|description|circumstance|what happened"
        r"|how.{0,15}(occur|happen)|basis.{0,10}of|details|injury|damage|loss", re.I)),
    ("employees",         re.compile(r"employee|officer|department|agency", re.I)),
    ("date_signed",       re.compile(r"^date$|date.{0,10}(signed|of.{0,5}(this.{0,5})?claim)|dated", re.I)),
]


def _fill_uploaded_claim_pdf(pdf_bytes: bytes, data: dict):
    """Best-effort fill of an uploaded local-jurisdiction claim form PDF.

    Matches the PDF's fillable text fields (name + tooltip) against common
    claim-form labels and fills what it can. Returns
    (filled_pdf_bytes, matched {label: value}, unmatched [labels]).
    Raises ValueError if the PDF has no fillable text fields.
    """
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import NameObject, BooleanObject, DictionaryObject

    reader = PdfReader(io.BytesIO(pdf_bytes))
    fields = reader.get_fields() or {}
    text_fields = {k: v for k, v in fields.items() if v.get("/FT") == "/Tx"}
    if not text_fields:
        raise ValueError(
            "This PDF has no fillable text fields — it is probably a scanned "
            "or flattened form. Print it and copy your answers from the "
            "generic claim form generated below."
        )

    values, matched, unmatched = {}, {}, []
    for fname, f in text_fields.items():
        label = str(f.get("/TU") or "").strip()
        haystack = f"{fname} {label}"
        for key, pattern in _LOCAL_CLAIM_PATTERNS:
            val = (data.get(key) or "").strip()
            if val and pattern.search(haystack):
                values[fname] = val
                matched[label or fname] = val
                break
        else:
            unmatched.append(label or fname)

    writer = PdfWriter()
    writer.append(reader)
    for page in writer.pages:
        writer.update_page_form_field_values(page, values)
    try:
        root = writer._root_object
        acro = None
        try:
            acro = root.get(NameObject("/AcroForm")) if hasattr(root, "get") else None
        except Exception:
            acro = None
        if acro is not None and hasattr(acro, "update"):
            acro.update({NameObject("/NeedAppearances"): BooleanObject(True)})
        else:
            root.update({NameObject("/AcroForm"): DictionaryObject(
                {NameObject("/NeedAppearances"): BooleanObject(True)}
            )})
    except Exception:
        pass

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue(), matched, unmatched


_DEFAULT_SUBPOENA_REQUESTS = [
    "All body-worn camera, dashboard camera, and other video or audio recordings from the sweep.",
    "All incident reports, field notes, supplemental reports, and emails by officers involved.",
    "All dispatch logs, radio transmissions, and 911/311 call recordings for the incident.",
    "All complaints, investigations, internal affairs files, and disciplinary records for involved officers.",
    "All policies, training materials, written directives, and encampment sweep protocols.",
    "All property seizure, storage, chain-of-custody, and disposal records.",
    "All internal communications, memos, and coordination records between police, DPW, and other City agencies.",
    "All surveillance camera and private video footage from the sweep location.",
    "All records authorizing, scheduling, or directing the sweep.",
]


@contextlib.contextmanager
def _quiet():
    with open(os.devnull, "w") as nul:
        old, sys.stdout = sys.stdout, nul
        try:
            yield
        finally:
            sys.stdout = old


def _generate_pdfs(case: dict) -> dict:
    """Fill all forms. Returns {label: bytes}. Raises ValueError on bad input."""
    validate_case(case)
    result = {}
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        with _quiet():
            fill_sc100(case, str(_TPL/"sc100.pdf"), str(tmp/"sc100.pdf"), _META_SC100)
            result["SC-100"] = (tmp/"sc100.pdf").read_bytes()

            fill_fw001(case, str(_TPL/"fw001.pdf"), str(tmp/"fw001.pdf"), _META_FW001)
            result["FW-001"] = (tmp/"fw001.pdf").read_bytes()

            fill_fw003(case, str(_TPL/"fw003.pdf"), str(tmp/"fw003.pdf"))
            result["FW-003"] = (tmp/"fw003.pdf").read_bytes()

            fill_sc112a(case, str(_TPL/"sc112a.pdf"), str(tmp/"sc112a.pdf"))
            result["SC-112A"] = (tmp/"sc112a.pdf").read_bytes()

            # SC-150 Request to Postpone Trial — only if postponement data present
            if has_postponement(case):
                try:
                    fill_sc150(case, str(_TPL/"sc150.pdf"), str(tmp/"sc150.pdf"))
                    result["SC-150"] = (tmp/"sc150.pdf").read_bytes()
                except Exception:
                    # Non-fatal: continue generating other forms
                    pass

            # SC-107 subpoena package: form with attachment boxes checked
            # plus Attachment 2a / 3 / 4 pages (only if subpoena info present)
            _sub = case.get("subpoena", {}) or {}
            if any(r for r in (_sub.get("requests") or []) if r) or (_sub.get("to") or "").strip():
                try:
                    fill_sc107(case, str(_TPL/"sc107.pdf"), str(tmp/"sc107.pdf"))
                    result["SC-107"] = (tmp/"sc107.pdf").read_bytes()
                except Exception:
                    # Non-fatal: continue generating other forms
                    pass

            # SC-100A: generate one form per additional defendant (if present)
            for i, ad in enumerate(case.get('additional_defendants', []) or [] , start=1):
                try:
                    outp = tmp/f"sc100a_defendant_{i}.pdf"
                    fill_sc100a_for_party(case, str(outp), ad, role='defendant')
                    result[f"SC-100A-DEF-{i}"] = outp.read_bytes()
                except Exception:
                    pass
    return result


def _make_zip(pdfs: dict, slug: str, flatten: bool = False) -> bytes:
    """Create a ZIP of the provided PDF bytes.

    By default (`flatten=False`) the function packages the exact in-memory
    PDF bytes (so the ZIP matches the individual PDF downloads shown in-UI).
    If `flatten=True` the function will try to rasterize + sanitize PDFs for
    maximum viewer compatibility.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for label, data in pdfs.items():
            out_bytes = data

            if flatten:
                # Try to rasterize/flatten PDF bytes so appearances are baked in for
                # viewers that ignore /NeedAppearances. If PyMuPDF isn't available,
                # fall back to the original bytes.
                try:
                    import fitz
                    # Open original bytes, render each page to an image-PDF, then
                    # combine pages into a new PDF bytes object.
                    doc = fitz.open(stream=data, filetype="pdf")
                    new = fitz.open()
                    mat = fitz.Matrix(2.0, 2.0)
                    for page in doc:
                        pix = page.get_pixmap(matrix=mat)
                        img_pdf = fitz.open("pdf", pix.tobytes("pdf"))
                        new.insert_pdf(img_pdf)
                    try:
                        out_bytes = new.write()
                    except Exception:
                        # Fallback: save to temp file then read
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmpf:
                            new.save(tmpf.name)
                            tmpname = tmpf.name
                        with open(tmpname, "rb") as tf:
                            out_bytes = tf.read()
                        os.unlink(tmpname)
                    new.close()
                    doc.close()
                except Exception:
                    out_bytes = data

                # Post-process to remove any remaining AcroForm and page /Annots
                # so ZIPs contain image-only PDFs that render reliably in viewers.
                try:
                    from pypdf import PdfReader, PdfWriter
                    from pypdf.generic import NameObject
                    rdr = PdfReader(io.BytesIO(out_bytes))
                    w = PdfWriter()
                    w.append(rdr)

                    # Remove page annotations
                    for p in w.pages:
                        try:
                            if '/Annots' in p:
                                p.pop('/Annots', None)
                        except Exception:
                            pass

                    # Remove AcroForm from root if present
                    try:
                        root = w._root_object
                        if NameObject('/AcroForm') in root:
                            try:
                                del root[NameObject('/AcroForm')]
                            except Exception:
                                try:
                                    root.pop(NameObject('/AcroForm'), None)
                                except Exception:
                                    pass
                    except Exception:
                        pass

                    bio = io.BytesIO()
                    w.write(bio)
                    out_bytes = bio.getvalue()
                except Exception:
                    # If post-processing fails, keep the original bytes
                    pass

            zf.writestr(f"{slug}_{label.lower().replace('-', '')}.pdf", out_bytes)
    return buf.getvalue()


def _show_downloads(pdfs: dict, slug: str, label: str = "") -> None:
    prefix = f"{label} — " if label else ""
    st.success(f"{prefix}Generated {len(pdfs)} forms.")
    st.download_button(
        "⬇️  Download All Forms (ZIP)",
        data=_make_zip(pdfs, slug, flatten=False),
        file_name=f"{slug}_forms.zip",
        mime="application/zip",
        type="primary",
        width="stretch",
        key=f"zip_{slug}",
    )
    # Optional: flattened ZIP for viewers that ignore form appearances
    st.download_button(
        "⬇️  Download All Forms (ZIP, flattened for compatibility)",
        data=_make_zip(pdfs, slug, flatten=True),
        file_name=f"{slug}_forms_flattened.zip",
        mime="application/zip",
        type="secondary",
        width="stretch",
        key=f"zip_flat_{slug}",
    )
    cols = st.columns(len(pdfs))
    for col, (lbl, data) in zip(cols, pdfs.items()):
        fname = f"{slug}_{lbl.lower().replace('-', '')}.pdf"
        with col:
            st.download_button(
                lbl, data=data, file_name=fname,
                mime="application/pdf", width="stretch",
                key=f"pdf_{slug}_{lbl}",
            )


# ─── Address / date parsing ───────────────────────────────────────────────────

_STATE_ZIP_RE  = re.compile(r"\b([A-Z]{2})\s+(\d{5}(?:-\d{4})?)\b")
_ZIP_RE        = re.compile(r"\b(\d{5})\b")
_KNOWN_CITIES  = re.compile(
    r"\b(Oakland|Emeryville|Richmond|Alameda|Hayward|Fremont|Berkeley|Newark|"
    r"Martinez|San Leandro|Walnut Creek|Pleasanton|Livermore|Castro Valley)\b",
    re.IGNORECASE,
)


def _parse_address(raw: str) -> dict:
    """Split a freeform US address into {street, city, state, zip}."""
    raw = str(raw).strip() if raw and not (isinstance(raw, float) and pd.isna(raw)) else ""
    if not raw:
        return {"street": "", "city": "", "state": "CA", "zip": ""}

    m = _STATE_ZIP_RE.search(raw)
    if m:
        state, zip_ = m.group(1), m.group(2)
        before = raw[: m.start()].rstrip(", ").strip()

        # If there's a comma, it separates street from city cleanly
        if "," in before:
            parts = [p.strip() for p in before.rsplit(",", 1)]
            street, city = parts[0], parts[1]
        else:
            # Find the last known city name in the text before the state
            city_match = None
            for cm in _KNOWN_CITIES.finditer(before):
                city_match = cm
            if city_match:
                street = before[: city_match.start()].strip()
                city = city_match.group(0).title()
            else:
                # Fallback: last word is city
                parts = before.rsplit(None, 1)
                street, city = (parts[0], parts[1]) if len(parts) == 2 else (before, "")

        return {"street": street or raw, "city": city, "state": state, "zip": zip_}

    # No STATE ZIP found — just extract ZIP if present
    z = _ZIP_RE.search(raw)
    return {"street": raw, "city": "", "state": "CA", "zip": z.group(1) if z else ""}


def _parse_date(raw) -> str:
    """Parse messy date strings → 'MM/DD/YYYY'. Returns '' on failure."""
    if not raw or (isinstance(raw, float) and pd.isna(raw)):
        return ""
    s = str(raw).strip()
    if not s:
        return ""
    # Handle ranges like "June 3 and 4 2025" → "June 3 2025"
    s = re.sub(r"\s+and\s+\d+", "", s)
    # Remove ordinal suffixes
    s = re.sub(r"(\d+)(st|nd|rd|th)\b", r"\1", s)
    # Remove slashes in "June 17/2026"
    s = re.sub(r"(\d+)/(\d{4})", r"\1 \2", s)
    try:
        return _dateutil.parse(s).strftime("%m/%d/%Y")
    except Exception:
        return s


# ─── Spreadsheet format detection ────────────────────────────────────────────

_INTAKE_COLS = {"Name", "Address", "Phone Number", "Location of Injury", "Date of Injury"}
_TEMPLATE_COL = "plaintiff_name"


def _detect_format(df: pd.DataFrame) -> str:
    cols = set(df.columns)
    if _INTAKE_COLS.issubset(cols):
        return "oakland_intake"
    if _TEMPLATE_COL in cols:
        return "template"
    return "unknown"


# ─── Oakland intake → case dict ───────────────────────────────────────────────

_CLAIM_REASON_TMPL = (
    "On {date}, the City of Oakland Department of Public Works conducted an encampment "
    "sweep at {location}, Oakland, CA. City employees seized and destroyed Plaintiff's "
    "personal property without adequate notice, without providing an opportunity to retrieve "
    "belongings, and without following the City's bag-and-tag policy. Plaintiff's property "
    "included clothing, sleeping equipment, tools, and personal documents. The City's actions "
    "violated the Fourth Amendment, Article I § 13 of the California Constitution, and the "
    "City's own policies."
)

_DECL_TMPL = (
    "I am the plaintiff in this action. On {date}, the City of Oakland Department of Public "
    "Works conducted an encampment sweep at {location}, Oakland, CA. Without adequate notice "
    "and without providing me an opportunity to retrieve my belongings, City employees seized "
    "and destroyed my personal property. My property included clothing, sleeping equipment, "
    "tools, and personal documents. At no time did City employees tag or store my property "
    "for later retrieval as required by the City's bag-and-tag policy. I have been unable to "
    "replace most of these items and have suffered ongoing hardship as a result. "
    "I declare under penalty of perjury under the laws of the State of California that the "
    "foregoing is true and correct."
)


def intake_row_to_case(row: pd.Series, defaults: dict) -> dict:
    def g(col, fallback=""):
        v = row.get(col, fallback)
        return fallback if (isinstance(v, float) and pd.isna(v)) else str(v).strip()

    addr = _parse_address(g("Address"))
    inc_date = _parse_date(g("Date of Injury"))
    location = g("Location of Injury")

    reason = defaults.get("claim_reason") or _CLAIM_REASON_TMPL.format(
        date=inc_date or "the date of the sweep",
        location=location or "the encampment",
    )
    decl = defaults.get("declaration") or _DECL_TMPL.format(
        date=inc_date or "the date of the sweep",
        location=location or "the encampment",
    )
    damages = defaults.get("damages_calculation") or (
        f"Property destroyed at encampment sweep at {location}. "
        f"Total estimated damages: ${defaults.get('claim_amount', '10000')}."
    )

    # Phone: take first number if multiple listed
    phone_raw = g("Phone Number")
    phone = re.split(r"[;,/]", phone_raw)[0].strip()

    return {
        "court": defaults.get("court", {}),
        "plaintiff": {
            "name":   g("Name"),
            "street": addr["street"],
            "city":   addr["city"],
            "state":  addr["state"],
            "zip":    addr["zip"],
            "phone":  phone,
            "email":  g("email"),
        },
        "defendant": defaults.get("defendant") or DEFENDANT_DEFAULTS["city_of_oakland"],
        "claim": {
            "amount":                defaults.get("claim_amount", "10000"),
            "reason":                reason,
            "incident_date":         inc_date,
            "damages_calculation":   damages,
            "govt_claim_filed_date": defaults.get("govt_claim_filed_date", ""),
            "items":                 [],
        },
        "filing": {
            "filing_date":      defaults.get("filing_date", ""),
            "demanded_payment": True,
        },
        "fee_waiver": {
            "basis":                  defaults.get("fw_basis", "5c"),
            "waive_option":           "all",
            "receives_medi_cal":      defaults.get("receives_medi_cal", False),
            "receives_snap":          defaults.get("receives_snap", False),
            "receives_calworks":      defaults.get("receives_calworks", False),
            "income_source_1":        defaults.get("income_source", ""),
            "income_amount_1":        defaults.get("income_amount", ""),
            "total_monthly_income":   defaults.get("total_income", ""),
            "expense_housing":        defaults.get("expense_housing", "0"),
            "expense_food":           defaults.get("expense_food", "0"),
            "expense_utilities":      "0",
            "expense_medical":        defaults.get("expense_medical", "0"),
            "expense_transport":      defaults.get("expense_transport", "0"),
            "total_monthly_expenses": defaults.get("total_expenses", ""),
        },
        "declaration": {
            "declarant_name": g("Name"),
            "content":        decl,
        },
    }


# ─── Generic template format → case dict ─────────────────────────────────────

def template_row_to_case(row: pd.Series) -> dict:
    def s(col, default=""):
        v = row.get(col, default)
        return default if (isinstance(v, float) and pd.isna(v)) else str(v).strip()

    def b(col, default=False):
        v = row.get(col, "")
        if isinstance(v, float) and pd.isna(v):
            return default
        return str(v).strip().upper() in ("TRUE", "YES", "1", "Y")

    items = []
    for i in range(1, 7):
        desc = s(f"item_{i}_desc")
        val  = s(f"item_{i}_value")
        if desc:
            items.append({"description": desc, "value": val})

    reason = s("claim_reason")
    return {
        "plaintiff": {
            "name":   s("plaintiff_name"),
            "street": s("plaintiff_street"),
            "city":   s("plaintiff_city"),
            "state":  s("plaintiff_state", "CA"),
            "zip":    s("plaintiff_zip"),
            "phone":  s("plaintiff_phone"),
            "email":  s("plaintiff_email"),
        },
        "defendant": DEFENDANT_DEFAULTS["city_of_oakland"],  # template format always uses spreadsheet data
        "claim": {
            "amount":                s("claim_amount"),
            "reason":                reason,
            "incident_date":         s("incident_date"),
            "damages_calculation":   s("damages_calculation") or reason,
            "govt_claim_filed_date": s("govt_claim_filed_date"),
            "items":                 items,
        },
        "filing": {
            "filing_date":      s("filing_date"),
            "demanded_payment": b("demanded_payment", True),
        },
        "fee_waiver": {
            "basis":                  s("fee_waiver_basis", "5c"),
            "waive_option":           "all",
            "receives_medi_cal":      b("receives_medi_cal"),
            "receives_snap":          b("receives_snap"),
            "receives_calworks":      b("receives_calworks"),
            "income_source_1":        s("income_source_1"),
            "income_amount_1":        s("income_amount_1"),
            "total_monthly_income":   s("total_monthly_income"),
            "expense_housing":        s("expense_housing", "0"),
            "expense_food":           s("expense_food", "0"),
            "expense_utilities":      "0",
            "expense_medical":        s("expense_medical", "0"),
            "expense_transport":      s("expense_transport", "0"),
            "total_monthly_expenses": s("total_monthly_expenses"),
        },
        "declaration": {
            "declarant_name": s("plaintiff_name"),
            "content":        s("declaration_content") or reason,
        },
        "subpoena": {
            "case_caption":          s("subpoena_case_caption"),
            "to":                    s("subpoena_to"),
            "custodian":             s("subpoena_custodian"),
            "service_location":      s("subpoena_service_location"),
            "requests": [
                s("subpoena_request_1"),
                s("subpoena_request_2"),
                s("subpoena_request_3"),
                s("subpoena_request_4"),
                s("subpoena_request_5"),
                s("subpoena_request_6"),
                s("subpoena_request_7"),
                s("subpoena_request_8"),
                s("subpoena_request_9"),
                s("subpoena_request_10"),
            ],
            "good_cause":  s("subpoena_good_cause"),
            "materiality": s("subpoena_materiality"),
        },
    }


# ─── CSV template for download ────────────────────────────────────────────────

_TEMPLATE_COLS = [
    ("plaintiff_name",          "Full legal name",                                True,  "Jane Doe"),
    ("plaintiff_street",        "Street / PO Box (c/o ... for unhoused)",         True,  "c/o 1234 Telegraph Ave"),
    ("incident_date",           "Date of sweep MM/DD/YYYY",                       True,  "05/12/2025"),
    ("claim_amount",            "Total claim dollars (max 12500)",                True,  "10000"),
    ("claim_reason",            "What happened — used on SC-100 and SC-150",      True,  "On May 12 2025 City of Oakland DPW..."),
    ("govt_claim_filed_date",   "Date govt tort claim filed with City Clerk",     True,  "08/15/2025"),
    ("filing_date",             "Date filing court papers MM/DD/YYYY",            True,  "09/15/2025"),
    ("total_monthly_income",    "Total monthly income $",                          True,  "400"),
    ("total_monthly_expenses",  "Total monthly expenses $",                        True,  "300"),
    ("plaintiff_city",          "Plaintiff's city",                               False, "Oakland"),
    ("plaintiff_state",         "State (default CA)",                             False, "CA"),
    ("plaintiff_zip",           "ZIP code",                                       False, "94609"),
    ("plaintiff_phone",         "Phone number",                                   False, "510-555-0100"),
    ("plaintiff_email",         "Email",                                          False, ""),
    ("damages_calculation",     "How damages were calculated (SC-100)",         False, "Clothing $500..."),
    ("income_source_1",         "Primary income source",                          False, "General Assistance"),
    ("income_amount_1",         "Primary income amount $",                         False, "400"),
    ("expense_food",            "Monthly food/supplies $",                         False, "200"),
    ("expense_medical",         "Monthly medical $",                               False, "50"),
    ("expense_transport",       "Monthly transport $",                             False, "50"),
    ("expense_housing",         "Monthly housing $",                               False, "0"),
    ("receives_medi_cal",       "Receives Medi-Cal? TRUE/FALSE",                  False, "TRUE"),
    ("fee_waiver_basis",        "Fee waiver basis: 5a 5b or 5c",                  False, "5c"),
    ("declaration_content",     "First-person declaration (optional)",            False, ""),

    # SC-107 subpoena helper fields
    ("subpoena_case_caption",   "SC-107 subpoena case caption",                   False, "Jane Doe v. City of Oakland"),
    ("subpoena_to",             "Subpoena recipient / agency",                    False, "Oakland Police Department"),
    ("subpoena_custodian",      "Custodian of records",                            False, "Records Division"),
    ("subpoena_service_location","Service address for subpoena",                   False, "1515 Clay St, Oakland CA"),
    ("subpoena_request_1",      "SC-107 request item 1",                          False, "All body-worn camera and officer dashboard footage from the sweep."),
    ("subpoena_request_2",      "SC-107 request item 2",                          False, "All police incident reports, notes, and supplemental reports related to the sweep."),
    ("subpoena_request_3",      "SC-107 request item 3",                          False, "All dispatch logs, radio transmissions, and 911/311 call recordings for the incident."),
    ("subpoena_request_4",      "SC-107 request item 4",                          False, "All complaints, investigations, and disciplinary records for involved officers."),
    ("subpoena_request_5",      "SC-107 request item 5",                          False, "All internal communications, emails, memos, and directives regarding encampment sweeps."),
    ("subpoena_request_6",      "SC-107 request item 6",                          False, "All policies, training materials, use-of-force guidelines, and homeless encampment protocols."),
    ("subpoena_request_7",      "SC-107 request item 7",                          False, "All property seizure, storage, chain-of-custody, and disposal records."),
    ("subpoena_request_8",      "SC-107 request item 8",                          False, "All surveillance camera and private video footage from the sweep location."),
    ("subpoena_request_9",      "SC-107 request item 9",                          False, "All records of coordination between police, DPW, and other City agencies."),
    ("subpoena_request_10",     "SC-107 request item 10",                         False, "All logs, schedules, and written directives authorizing the sweeps."),
    ("subpoena_good_cause",     "Attachment 3: why good cause exists (blank = default)",   False, ""),
    ("subpoena_materiality",    "Attachment 4: why records are material (blank = default)", False, ""),
    ("item_1_desc",             "Property item 1 description",                    False, "Tent and sleeping bag"),
    ("item_1_value",            "Property item 1 value $",                         False, "350"),
    ("item_2_desc",             "Property item 2 description",                    False, "Clothing"),
    ("item_2_value",            "Property item 2 value $",                         False, "500"),
]


def _csv_template_bytes() -> bytes:
    row = {c[0]: c[3] for c in _TEMPLATE_COLS}
    buf = io.StringIO()
    pd.DataFrame([row]).to_csv(buf, index=False)
    return buf.getvalue().encode()


_SC107_TEMPLATE_COLS = [
    ("subpoena_case_caption",    "SC-107 subpoena case caption",                   False, "Jane Doe v. City of Oakland"),
    ("subpoena_to",              "Subpoena recipient / agency",                    False, "Oakland Police Department"),
    ("subpoena_custodian",       "Custodian of records",                            False, "Records Division"),
    ("subpoena_service_location","Service address for subpoena",                   False, "1515 Clay St, Oakland CA"),
    ("subpoena_request_1",       "SC-107 request item 1",                          False, "All body-worn camera and officer dashboard footage from the sweep."),
    ("subpoena_request_2",       "SC-107 request item 2",                          False, "All police incident reports, notes, and supplemental reports related to the sweep."),
    ("subpoena_request_3",       "SC-107 request item 3",                          False, "All dispatch logs, radio transmissions, and 911/311 call recordings for the incident."),
    ("subpoena_request_4",       "SC-107 request item 4",                          False, "All complaints, investigations, and disciplinary records for involved officers."),
    ("subpoena_request_5",       "SC-107 request item 5",                          False, "All internal communications, emails, memos, and directives regarding encampment sweeps."),
    ("subpoena_request_6",       "SC-107 request item 6",                          False, "All policies, training materials, use-of-force guidelines, and homeless encampment protocols."),
    ("subpoena_request_7",       "SC-107 request item 7",                          False, "All property seizure, storage, chain-of-custody, and disposal records."),
    ("subpoena_request_8",       "SC-107 request item 8",                          False, "All surveillance camera and private video footage from the sweep location."),
    ("subpoena_request_9",       "SC-107 request item 9",                          False, "All records of coordination between police, DPW, and other City agencies."),
    ("subpoena_request_10",      "SC-107 request item 10",                         False, "All logs, schedules, and written directives authorizing the sweeps."),
    ("subpoena_good_cause",      "Attachment 3: why good cause exists (blank = default)",   False, ""),
    ("subpoena_materiality",     "Attachment 4: why records are material (blank = default)", False, ""),
]


def _sc107_csv_template_bytes() -> bytes:
    row = {c[0]: c[3] for c in _SC107_TEMPLATE_COLS}
    buf = io.StringIO()
    pd.DataFrame([row]).to_csv(buf, index=False)
    return buf.getvalue().encode()


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE
# ═══════════════════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="CA Small Claims Autofiller",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Match the main BHU site: serif headers, gold accents on black
st.markdown(
    """
    <style>
    h1, h2, h3 { font-family: Georgia, 'Times New Roman', serif !important; }
    h1 { color: #FFC700 !important; }
    div[data-testid="stMetricValue"] { color: #FFC700; }
    </style>
    """,
    unsafe_allow_html=True,
)

# ─── Officer sign-in (top right) & data portal ────────────────────────────────

from accounts import (
    add_user as _acct_add, load_users as _acct_load, verify_login as _acct_verify,
)


_MEDIA_TRACKER_URL = os.environ.get(
    "BHU_MEDIA_TRACKER_URL", "https://bhu-media-tracker.vercel.app"
)


def _admin_url() -> str:
    """URL of the full case tracker (admin.py), for the link inside the portal."""
    try:
        if "admin_url" in st.secrets:
            return str(st.secrets["admin_url"])
    except Exception:
        pass
    return os.environ.get("BHU_ADMIN_URL", "http://localhost:8502")


def _load_case_records() -> list:
    return [case for _, case in _load_case_files()]


def _portal_date(s):
    s = (str(s) if s is not None else "").strip()
    if not s:
        return None
    try:
        return _dateutil.parse(s, dayfirst=False).date()
    except Exception:
        return None


# ─── Case pipeline: the stages a claimant moves through ───────────────────────
# One source of truth for stage + defendant grouping, reused across the portal.
# Stage keys map onto the canonical STATUSES vocabulary from the case tracker.
_PIPELINE = [
    ("intake",     "Intake",        "📝"),
    ("govt_claim", "Govt claim",    "🏛️"),
    ("filed",      "Lawsuit filed", "⚖️"),
    ("trial_prep", "Trial prep",    "📎"),
    ("judgment",   "Judgment",      "🏁"),
]
_PIPELINE_INDEX = {k: i for i, (k, _l, _e) in enumerate(_PIPELINE)}


def _stage_key(c: dict) -> str:
    """Pipeline stage from the officer-set status — the single source of truth.
    STATUSES already encodes the pipeline, so the status alone decides the stage."""
    status = str((c.get("tracking") or {}).get("status", "Intake"))
    if status.startswith(("Resolved", "Closed")):
        return "judgment"
    if status == "Trial Scheduled":
        return "trial_prep"
    if status == "Lawsuit Filed (SC-100)":
        return "filed"
    if status in ("Govt Claim Filed", "Claim Rejected / 45 Days Passed"):
        return "govt_claim"
    return "intake"


def _stage_meta(c: dict):
    key = _stage_key(c)
    idx = _PIPELINE_INDEX[key]
    return key, idx, _PIPELINE[idx][1], _PIPELINE[idx][2]


def _case_stage(c: dict) -> str:
    """Compact label for list rows, e.g. '⚖️ Lawsuit filed'."""
    _k, _i, label, emoji = _stage_meta(c)
    return f"{emoji} {label}"


def _stage_bar_html(c: dict) -> str:
    """A 5-step pipeline 'symbol' with the current stage highlighted."""
    _k, idx, _label, _emoji = _stage_meta(c)
    parts = []
    for i, (_kk, lab, em) in enumerate(_PIPELINE):
        if i < idx:
            parts.append(f"<span style='color:#FFC700'>●&nbsp;{lab}</span>")
        elif i == idx:
            parts.append(
                f"<span style='color:#FFC700;font-weight:700'>{em}&nbsp;{lab}</span>"
            )
        else:
            parts.append(f"<span style='color:#6f6a55'>○&nbsp;{lab}</span>")
    return "<div style='font-size:0.9em'>" + " &nbsp;→&nbsp; ".join(parts) + "</div>"


def _defendant_key(c: dict) -> str:
    """Normalized grouping key so co-plaintiffs against the same entity cluster."""
    name = ((c.get("defendant") or {}).get("name") or "").strip().lower()
    name = re.sub(r"[^a-z0-9 ]+", " ", name)
    name = re.sub(r"^the\s+", "", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name or "unspecified"


def _defendant_label(c: dict) -> str:
    return ((c.get("defendant") or {}).get("name") or "").strip() or "Unspecified defendant"


def _grouped_records(records: list) -> list:
    """Records sorted so people suing the same defendant sit together,
    most-sued defendants first — so a whole group pulls out easily for trial."""
    counts = {}
    for c in records:
        counts[_defendant_key(c)] = counts.get(_defendant_key(c), 0) + 1
    return sorted(
        records,
        key=lambda c: (-counts[_defendant_key(c)], _defendant_key(c),
                       c.get("internal_case_number", "")),
    )


def _master_dataframe(records: list):
    """One row per claimant, grouped by defendant, with derived stage +
    defendant_group columns. Field columns come straight from the records,
    so new intake fields appear automatically."""
    recs = _grouped_records(records)
    flat = pd.json_normalize(recs, sep=".")
    flat.insert(0, "defendant_group", [_defendant_label(c) for c in recs])
    flat.insert(1, "stage", [_case_stage(c) for c in recs])

    def _cell(v):
        return json.dumps(v, ensure_ascii=False) if isinstance(v, (list, dict)) else v

    return flat.apply(lambda col: col.map(_cell)).fillna("")


def _case_alerts(c: dict):
    """(flag, detail) — same deadline rules as the case tracker."""
    from datetime import timedelta

    today = datetime.now().date()
    cl = c.get("claim") or {}
    status = (c.get("tracking") or {}).get("status", "Intake")

    incident = _portal_date(cl.get("incident_date"))
    if status == "Intake" and incident:
        left = (incident + timedelta(days=182) - today).days
        if left < 0:
            return "🔴", f"Govt claim window passed {-left}d ago"
        if left <= 30:
            return "🟠", f"{left}d left to file govt claim"

    claim_filed = _portal_date(cl.get("govt_claim_filed_date"))
    if status == "Govt Claim Filed" and claim_filed:
        over = (today - (claim_filed + timedelta(days=45))).days
        if over >= 0:
            return "🟢", "45 days passed — can file lawsuit now"
        if over >= -7:
            return "🟡", f"45-day mark in {-over}d"

    if not status.startswith(("Resolved", "Closed")):
        trial = _portal_date((c.get("lawsuit") or {}).get("trial_date"))
        if trial:
            days = (trial - today).days
            if days == 0:
                return "🔵", "Trial TODAY"
            if 0 < days <= 14:
                return "🔵", f"Trial in {days}d"
    return "", ""


_PORTAL_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=Playfair+Display:ital,wght@0,600;0,700;1,600&display=swap');
:root{
  --bhu-black:#080808; --bhu-off:#101010; --bhu-rule:#2F2F2F;
  --bhu-accent:#FFC700; --bhu-mid:#C8A900; --bhu-light:#F9E873; --bhu-text:#F5EFDC;
}
.stApp{ background:var(--bhu-black); }
html, body, .stMarkdown, p, span, div, label, input, textarea, button{
  font-family:'Inter', system-ui, sans-serif;
}
h1,h2,h3,h4,[data-testid="stHeading"]{
  font-family:'Playfair Display', Georgia, serif !important;
  color:var(--bhu-text) !important; letter-spacing:.01em;
}
[data-testid="stMetric"]{
  background:var(--bhu-off); border:1px solid var(--bhu-rule);
  border-radius:14px; padding:14px 18px;
}
[data-testid="stMetricValue"]{ font-family:'Playfair Display', Georgia, serif; color:var(--bhu-accent) !important; }
[data-testid="stMetricLabel"] p{
  text-transform:uppercase; letter-spacing:.14em; font-size:.72rem !important; color:var(--bhu-mid) !important;
}
button[data-baseweb="tab"]{ letter-spacing:.05em; }
button[data-baseweb="tab"][aria-selected="true"]{ color:var(--bhu-accent) !important; }
[data-baseweb="tab-highlight"]{ background-color:var(--bhu-accent) !important; }
[data-testid="stExpander"]{
  border:1px solid var(--bhu-rule) !important; border-radius:12px !important;
  background:var(--bhu-off) !important; margin-bottom:10px; overflow:hidden;
  transition:border-color .2s ease;
}
[data-testid="stExpander"]:hover{ border-color:var(--bhu-accent) !important; }
[data-testid="stExpander"] summary{ padding:10px 16px; font-weight:600; }
[data-testid="stExpander"] summary:hover{ color:var(--bhu-accent); }
[data-testid="stExpander"] h5{
  text-transform:uppercase; letter-spacing:.12em; font-size:.72rem !important;
  color:var(--bhu-mid) !important; font-family:'Inter', sans-serif !important; margin-bottom:.15rem;
}
blockquote{ border-left:2px solid var(--bhu-mid) !important; color:var(--bhu-light) !important; }
.stDownloadButton button, .stButton button{
  background:var(--bhu-accent) !important; color:#080808 !important;
  border:none !important; border-radius:999px !important;
  font-weight:700 !important; letter-spacing:.02em;
}
.stDownloadButton button:hover, .stButton button:hover{ background:var(--bhu-light) !important; color:#080808 !important; }
hr{ border-color:var(--bhu-rule) !important; }
</style>
"""


def _inject_portal_css() -> None:
    """Black & gold styling to match berkeleyhomelessunion.org — Playfair
    headings, Inter body, gold accents, carded panels."""
    st.markdown(_PORTAL_CSS, unsafe_allow_html=True)


def _render_admin_portal(user: str) -> None:
    """Signed-in officers see the collected data instead of the intake form."""
    _inject_portal_css()
    st.header("📋 Officer Data Portal")
    st.caption(
        f"Signed in as **{user}**. Everything the site has collected, one "
        "section per claimant. To change statuses, fix intake mistakes, or "
        "generate filing packets, open the "
        f"[full case tracker]({_admin_url()})."
    )

    records = _load_case_records()
    if not records:
        st.info(
            "No claimant data yet. A claimant's record appears here the "
            "moment they press **Save Progress** or **Generate Forms** on "
            "this site."
        )
        return

    m1, m2, m3 = st.columns(3)
    m1.metric("Claimants", len(records))
    m2.metric(
        "Active cases",
        sum(1 for c in records
            if not str((c.get("tracking") or {}).get("status", "Intake")).startswith(("Resolved", "Closed"))),
    )
    _total = 0.0
    for c in records:
        try:
            _total += float(str((c.get("claim") or {}).get("amount", "0")).replace("$", "").replace(",", "") or 0)
        except ValueError:
            pass
    m3.metric("Total claimed", f"${_total:,.0f}")
    st.divider()

    tab_today, tab_people, tab_csv, tab_media = st.tabs(
        ["🏠 Today", "👥 Members", "📊 Data (CSV)", "📰 Media Tracker"]
    )

    with tab_today:
        st.subheader("⚠️ Needs attention")
        flagged = []
        for c in records:
            fl, detail = _case_alerts(c)
            if fl:
                flagged.append((fl, detail, c))
        if not flagged:
            st.success(
                "Nothing urgent — no claim windows closing, no 45-day marks, "
                "no trials in the next two weeks."
            )
        for fl, detail, c in flagged:
            p = c.get("plaintiff") or {}
            r1, r2 = st.columns([4, 1])
            with r1:
                st.markdown(
                    f"{fl} **{p.get('name', '')}** "
                    f"({c.get('internal_case_number', '—')}) — {detail}"
                )
            with r2:
                if (p.get("phone") or "").strip():
                    st.markdown(f"[📞 {p['phone']}](tel:{p['phone']})")

        st.divider()
        st.subheader("📅 Upcoming trials")
        today_d = datetime.now().date()
        trials = sorted(
            ((_portal_date((c.get("lawsuit") or {}).get("trial_date")), c) for c in records),
            key=lambda x: x[0] or today_d,
        )
        trials = [(d_, c) for d_, c in trials if d_ and d_ >= today_d][:8]
        if trials:
            for d_, c in trials:
                p = c.get("plaintiff") or {}
                lw = c.get("lawsuit") or {}
                st.markdown(
                    f"**{d_:%a, %b %d}** — {p.get('name', '')} v. "
                    f"{(c.get('defendant') or {}).get('name', '—')} · "
                    f"Dept {lw.get('department') or '—'} · "
                    f"{c.get('case_number') or 'no case # yet'}"
                )
        else:
            st.caption("No trials scheduled.")

        st.divider()
        st.subheader("🆕 Recent activity")
        def _last_touch(c):
            return max(
                c.get("captured_at") or "",
                (c.get("tracking") or {}).get("updated_at") or "",
                c.get("forms_generated_at") or "",
            )
        for c in sorted(records, key=_last_touch, reverse=True)[:6]:
            p = c.get("plaintiff") or {}
            st.markdown(
                f"{_case_stage(c)} · **{p.get('name', '')}** "
                f"({c.get('internal_case_number', '—')}) · "
                f"last activity {(_last_touch(c) or '—')[:16].replace('T', ' ')}"
            )

    with tab_people:
        _grouped = _grouped_records(records)
        _seen_group = None
        for c in _grouped:
            _dk = _defendant_key(c)
            if _dk != _seen_group:
                _seen_group = _dk
                _peers = [x for x in _grouped if _defendant_key(x) == _dk]
                _hcol, _bcol = st.columns([3, 1])
                _hcol.markdown(
                    f"### ⚖️ {_defendant_label(c)} · {len(_peers)} "
                    f"claimant{'s' if len(_peers) != 1 else ''}"
                )
                _bcol.download_button(
                    "⬇️ Group CSV",
                    _master_dataframe(_peers).to_csv(index=False).encode(),
                    file_name=f"group_{_slug(_defendant_label(c))}.csv",
                    mime="text/csv",
                    key=f"grpcsv_{_dk}",
                    use_container_width=True,
                )
            p = c.get("plaintiff") or {}
            cl = c.get("claim") or {}
            d = c.get("defendant") or {}
            t = c.get("tracking") or {}
            lw = c.get("lawsuit") or {}
            label = (
                f"**{c.get('internal_case_number', '—')}** · {p.get('name', '')} "
                f"· {_case_stage(c)} · {t.get('status', 'Intake')} · ${cl.get('amount', '—')}"
            )
            with st.expander(label):
                st.markdown(_stage_bar_html(c), unsafe_allow_html=True)
                a, b = st.columns(2)
                with a:
                    st.markdown(
                        "##### Contact\n"
                        f"{p.get('name', '—')}  \n"
                        f"{p.get('street', '—')}, {p.get('city', '')} {p.get('state', '')} {p.get('zip', '')}  \n"
                        f"📞 {p.get('phone') or '—'} · ✉️ {p.get('email') or '—'}"
                    )
                    st.markdown(
                        "##### Claim\n"
                        f"**Against:** {d.get('name', '—')}  \n"
                        f"**Incident:** {cl.get('incident_date', '—')} · "
                        f"**Amount:** ${cl.get('amount', '—')}  \n"
                        f"**Govt claim filed:** {cl.get('govt_claim_filed_date') or '—'} · "
                        f"**Filing date:** {(c.get('filing') or {}).get('filing_date') or '—'}"
                    )
                    if cl.get("reason"):
                        st.markdown(f"> {cl['reason'][:400]}{'…' if len(cl.get('reason', '')) > 400 else ''}")
                    _items = [i for i in (cl.get("items") or []) if (i.get("description") or "").strip()]
                    if _items:
                        st.markdown("##### Itemized property")
                        st.table(pd.DataFrame(_items))
                with b:
                    st.markdown(
                        "##### Case activity\n"
                        f"**Status:** {t.get('status', 'Intake')}  \n"
                        f"**First captured:** {(c.get('captured_at') or '—')[:16]}  \n"
                        f"**Last update:** {(t.get('updated_at') or '—')[:16]}"
                    )
                    for h in reversed((t.get("history") or [])[-10:]):
                        st.caption(f"• {h.get('at', '')[:16]} · {h.get('officer') or '—'} · {h.get('change', '')}")
                    if lw:
                        st.markdown(
                            "##### Lawsuit\n"
                            f"**Court case #:** {c.get('case_number') or '—'} · "
                            f"**Filed:** {lw.get('filed_on') or '—'}  \n"
                            f"**Trial:** {lw.get('trial_date') or '—'} · "
                            f"**Outcome:** {lw.get('outcome') or 'Pending'} · "
                            f"**Judgment:** ${lw.get('judgment_amount') or '—'}"
                        )
                    _sub = c.get("subpoena") or {}
                    if (_sub.get("to") or "").strip() or any(_sub.get("requests") or []):
                        st.markdown(
                            "##### Subpoena\n"
                            f"**To:** {_sub.get('to') or '—'} · "
                            f"**{len([r for r in (_sub.get('requests') or []) if r])}** record request(s)"
                        )
                st.download_button(
                    "⬇️ Full record (JSON)",
                    data=json.dumps(c, indent=2).encode(),
                    file_name=f"{c.get('internal_case_number', 'case')}_record.json",
                    mime="application/json",
                    key=f"portal_json_{c.get('internal_case_number', '')}_{p.get('name', '')}",
                )

    with tab_csv:
        st.caption(
            "One row per claimant, grouped so people suing the same defendant "
            "sit together. Columns come straight from the intake records, so "
            "new form fields show up here automatically."
        )
        master = _master_dataframe(records)
        st.dataframe(master, use_container_width=True, height=420)
        st.download_button(
            "⬇️ Master CSV (all claimants, grouped by defendant)",
            data=master.to_csv(index=False).encode(),
            file_name=f"bhu_master_claims_{datetime.now():%Y%m%d}.csv",
            mime="text/csv",
            use_container_width=True,
        )

    with tab_media:
        _mc1, _mc2 = st.columns([4, 1])
        with _mc1:
            st.caption(
                "BHU Media Coverage Tracker — homelessness & street vendor "
                "coverage in Berkeley/Oakland. Updates automatically every "
                "morning."
            )
        with _mc2:
            st.link_button("↗ Full screen", _MEDIA_TRACKER_URL, use_container_width=True)
        try:
            from streamlit.components.v1 import iframe as _iframe
            _iframe(_MEDIA_TRACKER_URL, height=1200, scrolling=True)
        except Exception:
            st.info(f"Couldn't embed the tracker — open it directly: {_MEDIA_TRACKER_URL}")


_pop = st.popover if hasattr(st, "popover") else (lambda label, **_k: st.expander(label))
_title_l, _title_r = st.columns([5, 1])
with _title_l:
    st.title("California Encampment — Small Claims Autofiller")
with _title_r:
    _signed_in = st.session_state.get("bhu_admin_user")
    if _signed_in:
        with _pop(f"👤 {_signed_in}", use_container_width=True):
            if st.button("Sign out", use_container_width=True, key="portal_signout"):
                st.session_state.pop("bhu_admin_user", None)
                st.rerun()
    else:
        with _pop("🔐 Sign In", use_container_width=True):
            _users = _acct_load()
            if not _users:
                st.caption("No officer accounts yet — create the admin account:")
                nu = st.text_input("Admin username", key="portal_new_user")
                np1 = st.text_input("Password (min 8 chars)", type="password", key="portal_new_pw1")
                np2 = st.text_input("Confirm password", type="password", key="portal_new_pw2")
                if st.button("Create admin account", key="portal_create", use_container_width=True):
                    if np1 != np2:
                        st.error("Passwords don't match.")
                    else:
                        err = _acct_add(_users, nu, np1)
                        if err:
                            st.error(err)
                        else:
                            st.session_state["bhu_admin_user"] = nu.strip().lower()
                            st.rerun()
            else:
                lu = st.text_input("Username", key="portal_login_user")
                lp = st.text_input("Password", type="password", key="portal_login_pw")
                if st.button("Sign in", key="portal_login", use_container_width=True):
                    if _acct_verify(_users, lu, lp):
                        st.session_state["bhu_admin_user"] = lu.strip().lower()
                        st.rerun()
                    else:
                        st.error("Wrong username or password.")

if st.session_state.get("bhu_admin_user"):
    _render_admin_portal(st.session_state["bhu_admin_user"])
    st.stop()

st.caption(
    "Generates a government claim form plus SC-100, SC-100A, FW-001, FW-003, "
    "SC-112A, and SC-150 for encampment property destruction cases in "
    "any California county — from the initial government claim, to filing the "
    "lawsuit, to preparing for trial."
)


# ─── Court selector widget (reused in both tabs) ─────────────────────────────

def _court_selector(key_prefix: str, default_county: str = "Alameda") -> dict:
    """Render county + courthouse dropdowns. Returns a court dict for case["court"]."""
    col_county, col_house = st.columns([1, 2])
    with col_county:
        county = st.selectbox(
            "County *",
            ALL_COUNTIES,
            index=ALL_COUNTIES.index(default_county) if default_county in ALL_COUNTIES else 0,
            key=f"{key_prefix}_county",
        )
    houses = courthouses_for_county(county)
    house_labels = [f"{h['city']} — {h['name']}" for h in houses]
    with col_house:
        house_idx = st.selectbox(
            "Courthouse *",
            range(len(house_labels)),
            format_func=lambda i: house_labels[i],
            key=f"{key_prefix}_house",
        )
    chosen = houses[house_idx]
    return {
        "county":  county,
        "name":    chosen["name"],
        "address": chosen["address"],
        "city":    chosen["city"],
        "zip":     chosen["zip"],
    }

_CUSTOM_DEFENDANT = "✏️ No prefill — enter any defendant (person, county, unincorporated area…)"


def _defendant_block(key_prefix: str, def_id: int, is_primary: bool) -> dict:
    """One defendant entry: fully editable fields for suing anyone — a
    person, business, county agency, or city. An optional dropdown prefills
    the fields from the California municipality database, but nothing is
    ever locked to a city.

    The primary defendant (SC-100) also gets agent-for-service fields;
    additional defendants (SC-100A) get phone / mailing / job-title fields.
    """
    city_options = [_CUSTOM_DEFENDANT] + ALL_CITIES
    selected = st.selectbox(
        "Prefill from a California city (optional)",
        city_options,
        index=0,
        key=f"{key_prefix}_def{def_id}_city_sel",
        help=(
            "You can sue anyone — type the defendant's information directly "
            "in the fields below (a person, business, county agency, or an "
            "unincorporated county area). If you're suing a California city, "
            "picking it here just prefills the address and agent-for-service "
            "fields; everything stays editable."
        ),
    )
    if selected == _CUSTOM_DEFENDANT:
        d = {
            "name": "", "address": "", "city": "", "state": "CA", "zip": "",
            "agent_name": "", "agent_title": "",
            "agent_address": "", "agent_city": "", "agent_state": "CA", "agent_zip": "",
        }
        kp = f"{key_prefix}_def{def_id}_custom"
    else:
        d = defendant_info(selected)
        # Selection is part of the key so switching cities refreshes the defaults
        kp = f"{key_prefix}_def{def_id}_{selected}"

    c1, c2 = st.columns(2)
    with c1:
        name_v   = st.text_input("Defendant Name *", value=d["name"], key=f"{kp}_name")
        street_v = st.text_input("Street Address", value=d["address"], key=f"{kp}_street")
        city_v   = st.text_input("City", value=d["city"], key=f"{kp}_city")
        s1, s2 = st.columns(2)
        with s1:
            state_v = st.text_input("State", value=d.get("state", "CA"), key=f"{kp}_state")
        with s2:
            zip_v = st.text_input("ZIP", value=d["zip"], key=f"{kp}_zip")

    out = {
        "name":    name_v.strip(),
        "address": street_v.strip(),
        "street":  street_v.strip(),
        "city":    city_v.strip(),
        "state":   state_v.strip() or "CA",
        "zip":     zip_v.strip(),
    }

    with c2:
        if is_primary:
            out["agent_name"]    = st.text_input(
                "Agent for Service (Name)", value=d["agent_name"], key=f"{kp}_agent_name",
                help="Who accepts legal papers for the defendant. For a city this is "
                     "usually the City Clerk; leave blank when suing an individual.",
            ).strip()
            out["agent_title"]   = st.text_input("Agent Title", value=d["agent_title"], key=f"{kp}_agent_title").strip()
            out["agent_address"] = st.text_input("Agent Street", value=d["agent_address"], key=f"{kp}_agent_street").strip()
            out["agent_city"]    = st.text_input("Agent City", value=d["agent_city"], key=f"{kp}_agent_city").strip()
            out["agent_state"]   = "CA"
            out["agent_zip"]     = st.text_input("Agent ZIP", value=d["agent_zip"], key=f"{kp}_agent_zip").strip()
        else:
            out["phone"]     = st.text_input("Phone", key=f"{kp}_phone").strip()
            out["mailing"]   = st.text_input("Mailing Address (if different)", key=f"{kp}_mailing").strip()
            out["job_title"] = st.text_input("Job Title (if known)", key=f"{kp}_job_title").strip()

    if out["name"]:
        st.caption(f"**{out['name']}** · {out['address']}, {out['city']}, {out['state']} {out['zip']}")
    return out


tab_manual, tab_sheet = st.tabs(["📝 Manual Entry", "📊 Spreadsheet Import"])


# ══════════════════════════════════════════════════════
# TAB 1 — MANUAL ENTRY
# ══════════════════════════════════════════════════════

with tab_manual:
    # ════════════════════════════════════════════════════
    # STEP 1 — GOVERNMENT CLAIM (file this first)
    # ════════════════════════════════════════════════════
    st.header("Step 1 — Government Claim")
    st.caption(
        "Before you can sue a California city or public entity for property "
        "destroyed in a sweep, you must first file a government tort claim "
        "(Gov. Code §§ 905, 910) with that entity — generally within six months "
        "of the incident. Fill in your information and the incident below, "
        "then either upload your jurisdiction's own claim form (PDF) to "
        "auto-fill it, or generate a generic claim form to file with the "
        "City Clerk. The entity being claimed against comes from the "
        "Defendant section in Step 2."
    )

    # ── Plaintiff ──────────────────────────────────────────────────────
    st.subheader("Plaintiff")
    c1, c2 = st.columns(2)
    with c1:
        name   = st.text_input("Full Legal Name *", placeholder="Jane Doe")
        street = st.text_input(
            "Street / Mailing Address",
            placeholder="c/o 1234 Telegraph Ave  (use c/o for unhoused clients)",
        )
        phone  = st.text_input("Phone", placeholder="510-555-0100")
    with c2:
        city = st.text_input("City", placeholder="Your city")
        cs1, cs2 = st.columns(2)
        with cs1:
            state = st.text_input("State", value="CA")
        with cs2:
            zip_  = st.text_input("ZIP", placeholder="94609")
        email = st.text_input("Email (optional)", placeholder="")

    # ── Incident & Claim (shared by the claim form and the lawsuit) ────
    st.divider()
    st.subheader("Incident & Claim")
    c1, c2 = st.columns(2)
    with c1:
        _date_range = st.checkbox("Date range / multi-day incident", key="manual_date_range")
        if _date_range:
            _dc1, _dc2 = st.columns(2)
            with _dc1:
                _date_start = st.text_input("Start Date *", placeholder="MM/DD/YYYY", key="manual_date_start")
            with _dc2:
                _date_end = st.text_input("End Date *", placeholder="MM/DD/YYYY", key="manual_date_end")
            incident_date = _date_start.strip()
            if _date_end.strip():
                incident_date = f"{incident_date} – {_date_end.strip()}" if incident_date else _date_end.strip()
        else:
            incident_date = st.text_input("Date of Sweep *", placeholder="MM/DD/YYYY", key="manual_date_single")
        incident_location = st.text_input(
            "Location of Sweep",
            placeholder="E.g. E 12th St & 16th Ave underpass, Oakland",
        )
    with c2:
        claim_amount = st.text_input("Claim Amount ($) *", placeholder="10000")
        involved_employees = st.text_input(
            "City employees or agencies involved (if known)",
            placeholder="DPW crew, police officers, contractor…",
        )

    claim_reason = st.text_area(
        "Brief summary of what happened (used on the claim form and SC-100)",
        placeholder=(
            "On [date], the City of Oakland DPW conducted an encampment sweep "
            "at [location] and destroyed Plaintiff's personal property…"
        ),
        height=120,
    )

    # ── Itemized property (used on the claim form, declaration, SC-100) ─
    st.divider()
    st.subheader("Itemized Property")
    st.caption("List each item that was destroyed, its estimated value, and its condition before the loss.")
    items_df = st.data_editor(
        pd.DataFrame({
            "Description": ["", "", ""],
            "Value ($)": ["", "", ""],
            "Condition": ["New", "Good", "Fair"],
        }),
        num_rows="dynamic",
        use_container_width=True,
        column_config={
            "Description": st.column_config.TextColumn(width="large"),
            "Value ($)": st.column_config.TextColumn(width="small"),
            "Condition": st.column_config.SelectboxColumn(
                width="small",
                options=["New", "Good", "Fair", "Salvage"],
            ),
        },
        hide_index=True,
    )

    def _items_from_editor() -> list:
        out = []
        for _, r in items_df.iterrows():
            description = str(r.get("Description", "")).strip()
            if not description:
                continue
            out.append({
                "description": description,
                "value": str(r.get("Value ($)", "")).strip(),
                "condition": str(r.get("Condition", "")).strip() or "Unknown",
            })
        return out

    # ── Claim data shared by the uploaded-PDF and generic paths ────────
    def _collect_claim_data() -> dict:
        # Defendant widgets live in Step 2; the primary defendant's values
        # are published to session state each run.
        _gc_defs = st.session_state.get("manual_defendants_data") or [{}]
        _gc_def = _gc_defs[0]
        _clerk_addr = ", ".join(part for part in [
            (_gc_def.get("agent_address") or "").strip(),
            (_gc_def.get("agent_city") or "").strip(),
            f"CA {(_gc_def.get('agent_zip') or '').strip()}".strip(),
        ] if part)
        return {
            "entity":            _gc_def.get("name", ""),
            "clerk_address":     _clerk_addr,
            "claimant_name":     name.strip(),
            "claimant_address":  ", ".join(p for p in [
                street.strip(), city.strip(),
                f"{state.strip()} {zip_.strip()}".strip(),
            ] if p),
            "claimant_phone":    phone.strip(),
            "claimant_email":    email.strip(),
            "incident_date":     incident_date.strip(),
            "incident_location": incident_location.strip(),
            "description":       claim_reason.strip(),
            "employees":         involved_employees.strip(),
            "amount":            claim_amount.strip(),
            "items":             _items_from_editor(),
        }

    # ── Option A: your jurisdiction's own claim form (PDF upload) ──────
    st.divider()
    st.markdown("**Option A — Use your jurisdiction's own claim form**")
    st.caption(
        "Many cities and counties require claims to be submitted on their "
        "own form (usually available on the city clerk's or county's "
        "website). Upload that form as a PDF and your answers above will "
        "be filled into it automatically where possible."
    )
    local_claim_pdf = st.file_uploader(
        "Upload your local jurisdiction's claim form (PDF)",
        type=["pdf"],
        key="govt_local_pdf",
    )

    if local_claim_pdf is not None:
        if st.button(
            "Auto-Fill Uploaded Claim Form", type="primary",
            use_container_width=True, key="gen_local_claim",
        ):
            try:
                filled, matched, unmatched = _fill_uploaded_claim_pdf(
                    local_claim_pdf.getvalue(), _collect_claim_data()
                )
                st.session_state["local_claim_bytes"] = filled
                st.session_state["local_claim_matched"] = matched
                st.session_state["local_claim_unmatched"] = unmatched
                st.session_state["local_claim_name"] = (
                    f"{_slug(name.strip() or 'claim')}_"
                    f"{_slug(Path(local_claim_pdf.name).stem)}_filled.pdf"
                )
            except ValueError as e:
                st.session_state.pop("local_claim_bytes", None)
                st.warning(str(e))
            except Exception as e:
                st.session_state.pop("local_claim_bytes", None)
                st.error(f"Could not fill the uploaded form: {e}")

        if st.session_state.get("local_claim_bytes"):
            _n_matched = len(st.session_state.get("local_claim_matched", {}))
            _unmatched = st.session_state.get("local_claim_unmatched", [])
            st.success(f"Auto-filled {_n_matched} field(s) on the uploaded form.")
            st.download_button(
                "⬇️ Download Filled Local Claim Form (PDF)",
                data=st.session_state["local_claim_bytes"],
                file_name=st.session_state.get("local_claim_name", "local_claim_filled.pdf"),
                mime="application/pdf",
                use_container_width=True,
                key="dl_local_claim",
            )
            if st.session_state.get("local_claim_matched"):
                with st.expander("Fields that were auto-filled"):
                    for lbl, val in st.session_state["local_claim_matched"].items():
                        st.markdown(f"- **{lbl}** → {val[:120]}")
            if _unmatched:
                with st.expander(
                    f"{len(_unmatched)} field(s) left blank — complete by hand"
                ):
                    for lbl in _unmatched:
                        st.markdown(f"- {lbl}")
            st.caption(
                "⚠️ Automatic matching is best-effort — **review every page** "
                "before signing and filing. Fields the matcher could not "
                "recognize are left blank."
            )

    # ── Option B: generic Gov. Code claim form (Word) ───────────────────
    st.markdown("**Option B — Generic claim form (Word)**")
    st.caption(
        "If your jurisdiction does not require its own form (or you can't "
        "get it), generate a generic claim that satisfies Gov. Code "
        "§§ 905/910."
    )
    if st.button(
        "Generate Generic Claim Form (Word)",
        use_container_width=True, key="gen_govt_claim",
    ):
        try:
            st.session_state["govt_claim_bytes"] = _build_govt_claim_docx(
                _collect_claim_data()
            )
            st.session_state["govt_claim_name"] = (
                f"{_slug(name.strip() or 'claim')}_government_claim.docx"
            )
        except Exception as e:
            st.session_state.pop("govt_claim_bytes", None)
            st.error(f"Could not generate the claim form: {e}")
    if st.session_state.get("govt_claim_bytes"):
        st.download_button(
            "⬇️ Download Government Claim (Word)",
            data=st.session_state["govt_claim_bytes"],
            file_name=st.session_state.get("govt_claim_name", "government_claim.docx"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_govt_claim",
        )
    st.caption(
        "Print, sign, and file the claim with the City Clerk (the agent for "
        "service shown in the Step 2 defendant section). The city generally has "
        "**45 days** to respond. Once the claim is rejected — or 45 days "
        "pass — move to Step 2."
    )

    # ════════════════════════════════════════════════════
    # STEP 2 — FILE THE SMALL CLAIMS LAWSUIT
    # ════════════════════════════════════════════════════
    st.divider()
    st.header("Step 2 — File the Small Claims Lawsuit")
    st.caption(
        "After the claim is rejected (or 45 days pass with no response), file in "
        "small claims court. Name the defendant(s), pick the court, set your "
        "dates, itemize the property, and generate the filing packet: SC-100 "
        "(+ SC-100A), FW-001, FW-003, and SC-112A."
    )

    # ── Defendants (dynamic list; extras go on SC-100A) ────────────────
    if "manual_def_ids" not in st.session_state:
        st.session_state["manual_def_ids"] = [0]
        st.session_state["manual_def_next"] = 1

    hdr_l, hdr_r = st.columns([0.92, 0.08])
    with hdr_l:
        st.subheader("Defendant")
    with hdr_r:
        if st.button("➕", key="manual_add_def", help="Add another defendant (listed on form SC-100A)"):
            st.session_state["manual_def_ids"].append(st.session_state["manual_def_next"])
            st.session_state["manual_def_next"] += 1
            st.rerun()

    manual_defendants = []
    for pos, def_id in enumerate(st.session_state["manual_def_ids"]):
        is_primary = pos == 0
        if is_primary:
            if len(st.session_state["manual_def_ids"]) > 1:
                st.markdown("**Defendant 1** · named on SC-100")
        else:
            rc1, rc2 = st.columns([0.92, 0.08])
            with rc1:
                st.markdown(f"**Defendant {pos + 1}** · on attached SC-100A")
            with rc2:
                if st.button("✕", key=f"manual_rm_def{def_id}", help="Remove this defendant"):
                    st.session_state["manual_def_ids"].remove(def_id)
                    st.rerun()
        manual_defendants.append(_defendant_block("manual", def_id, is_primary))
    # Step 1's government claim reads the primary defendant from here
    st.session_state["manual_defendants_data"] = manual_defendants
    st.divider()

    # Court selector lives outside the form so county → courthouse cascade works
    st.subheader("Filing Court")
    manual_court = _court_selector("manual")
    st.caption(
        f"Court: **Superior Court of California, County of {manual_court['county']}** · "
        f"{manual_court['address']}, {manual_court['city']}, CA {manual_court['zip']}"
    )

    fd1, fd2 = st.columns(2)
    with fd1:
        filing_date = st.text_input("Filing Date *", placeholder="MM/DD/YYYY")
    with fd2:
        govt_claim_date = st.text_input(
            "Govt Claim Filed with City Clerk *", placeholder="MM/DD/YYYY",
            help="The date you filed (or will file) the Step 1 government claim.",
        )

    declaration_text_input = st.text_area(
        "Write your declaration in your own words",
        placeholder="Start typing what happened. For example: I was present when the City took my belongings, I was not given notice, and I saw them throw away my property.",
        height=180,
    )

    if st.button("Generate declaration", use_container_width=True):
        items = _items_from_editor()

        declaration_text = _build_guided_declaration(
            declaration_text_input,
            {
                "incident_date": incident_date.strip(),
                "claim_amount": claim_amount.strip(),
                "items": items,
            },
        )
        st.session_state["declaration_text"] = declaration_text

    declaration_text = st.session_state.get("declaration_text", "")
    st.text_area(
        "Declaration draft",
        value=declaration_text or "Press the button above to generate a court-style declaration.",
        height=260,
    )
    damages_calc = st.text_area(
        "How Damages Are Calculated",
        placeholder=(
            "Itemize property value + emotional distress. "
            "Leave blank to auto-fill from description above."
        ),
        height=80,
    )

    # ── Fee Waiver ─────────────────────────────────────────────────────
    st.divider()
    st.subheader("Fee Waiver")
    c1, c2, c3 = st.columns(3)
    with c1:
        fw_basis = st.radio(
            "Basis",
            ["5c — Cannot afford fees", "5a — Public benefits", "5b — Income below threshold"],
            help="5c is correct for most encampment sweep clients.",
        )
        st.markdown("**Public benefits:**")
        recv_medi_cal = st.checkbox("Medi-Cal")
        recv_snap     = st.checkbox("CalFresh / SNAP")
        recv_calworks = st.checkbox("CalWORKS")
    with c2:
        income_source = st.text_input("Income Source", placeholder="General Assistance, SSI…")
        income_amount = st.text_input("Monthly Income ($)", placeholder="400")
        total_income  = st.text_input("Total Monthly Income ($)", placeholder="400")
    with c3:
        exp_food      = st.text_input("Food / Supplies ($)", value="0")
        exp_medical   = st.text_input("Medical / Dental ($)", value="0")
        exp_transport = st.text_input("Transportation ($)", value="0")
        exp_housing   = st.text_input("Housing ($)", value="0")
        total_expenses = st.text_input("Total Monthly Expenses ($)", placeholder="300")

    _gen_col, _save_col = st.columns([3, 1])
    with _gen_col:
        submitted = st.button(
            "Generate Forms", type="primary", use_container_width=True
        )
    with _save_col:
        save_progress = st.button(
            "💾 Save Progress",
            use_container_width=True,
            help="Save everything entered so far under your case number — no "
                 "forms are generated. You can come back later, and officers "
                 "can review and correct it in the case tracker.",
        )

    st.download_button(
        "⬇️ Download declaration as Word document",
        data=_build_declaration_docx(declaration_text),
        file_name="guided_declaration.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        disabled=not declaration_text.strip(),
    )

    # ── Handle submission / save-progress ──────────────────────────────────
    if submitted or save_progress:
        items = _items_from_editor()
        basis_code = fw_basis.split(" — ")[0].strip()

        _defendant = manual_defendants[0]

        case = {
            "court": manual_court,
            "plaintiff": {
                "name":   name.strip(),
                "street": street.strip(),
                "city":   city.strip(),
                "state":  state.strip(),
                "zip":    zip_.strip(),
                "phone":  phone.strip(),
                "email":  email.strip(),
            },
            "defendant": _defendant,
            "claim": {
                "amount":                claim_amount.strip(),
                "reason":                claim_reason.strip(),
                "incident_date":         incident_date.strip(),
                "damages_calculation":   damages_calc.strip() or claim_reason.strip(),
                "govt_claim_filed_date": govt_claim_date.strip(),
                "items":                 items,
            },
            "filing": {
                "filing_date":      filing_date.strip(),
                "demanded_payment": True,
            },
            "fee_waiver": {
                "basis":                  basis_code,
                "waive_option":           "all",
                "receives_medi_cal":      recv_medi_cal,
                "receives_snap":          recv_snap,
                "receives_calworks":      recv_calworks,
                "income_source_1":        income_source.strip(),
                "income_amount_1":        income_amount.strip(),
                "total_monthly_income":   total_income.strip(),
                "expense_housing":        exp_housing.strip(),
                "expense_food":           exp_food.strip(),
                "expense_utilities":      "0",
                "expense_medical":        exp_medical.strip(),
                "expense_transport":      exp_transport.strip(),
                "total_monthly_expenses": total_expenses.strip(),
            },
            "declaration": {
                "declarant_name": name.strip(),
                "content":        declaration_text.strip() or claim_reason.strip(),
            },
            "subpoena": {
                "case_caption":     f"{name.strip()} v. {_defendant.get('name') or 'City of Oakland'}",
                "to":               st.session_state.get("sub_recipient_name", "").strip(),
                "custodian":        st.session_state.get("sub_recipient_custodian", "").strip(),
                "service_location": st.session_state.get("sub_recipient_service", "").strip(),
                "requests": (
                    [r for _i, r in enumerate(_DEFAULT_SUBPOENA_REQUESTS)
                     if st.session_state.get(f"sub_req_{_i}", True)]
                    + [line.strip()
                       for line in st.session_state.get("sub_extra_requests", "").splitlines()
                       if line.strip()]
                )[:10],
                "good_cause":  st.session_state.get("sub_good_cause", "").strip(),
                "materiality": st.session_state.get("sub_materiality", "").strip(),
            },
        }

        # Attach additional defendants (each generates a filled SC-100A)
        case['additional_defendants'] = [
            dd for dd in manual_defendants[1:] if dd.get('name')
        ]

        # Capture the full intake under an internal case number before
        # generating anything, so the information is kept even if a form fails.
        try:
            _capture_case_record(case)
            st.info(
                f"Internal case number: **{case['internal_case_number']}** — "
                "the full intake was saved to the cases folder."
            )
        except Exception as e:
            case.setdefault("internal_case_number", f"{datetime.now():%Y%m%d}-XX")
            st.warning(f"Could not save the intake record: {e}")

        if save_progress and not submitted:
            st.success(
                "Progress saved. Come back anytime — re-saving on the same "
                "day updates the same record — and officers can review and "
                "correct everything in the case tracker."
            )
        else:
            try:
                pdfs = _generate_pdfs(case)
                # Record that this member has generated their forms (stage
                # tracking in the officer portal) and refresh the record.
                case["forms_generated_at"] = datetime.now().isoformat(timespec="seconds")
                try:
                    _capture_case_record(case)
                except Exception:
                    pass
                _show_downloads(pdfs, _slug(name.strip()))
                st.download_button(
                    "💾  Save Case Data (JSON)",
                    data=json.dumps(case, indent=2).encode(),
                    file_name=f"{case['internal_case_number']}_{_slug(name.strip())}_case.json",
                    mime="application/json",
                )
            except ValueError as e:
                st.error(str(e))
            except Exception as e:
                st.error(f"Unexpected error: {e}")


    # ════════════════════════════════════════════════════
    # STEP 3 — PREPARE FOR TRIAL
    # ════════════════════════════════════════════════════
    st.divider()
    st.header("Step 3 — Prepare for Trial")
    st.caption(
        "Once your case is filed, gather your evidence. Use the subpoena "
        "below to make the city or other agencies produce records — footage, "
        "reports, policies — before your hearing. If you'll need someone to "
        "help you present your case at trial, use the SC-109 section below "
        "the subpoena. If you can't make your trial date, use the SC-150 "
        "section to ask the court to postpone the trial."
    )

    # ── Subpoena (SC-107) — checkboxes only + typed attachments ─────────
    st.divider()
    st.subheader("Subpoena Request (SC-107)")
    st.caption(
        "Generates the SC-107 with only the \"Continued on Attachment "
        "2a / 3 / 4\" boxes checked — the form's own text fields are left "
        "blank for you to complete by hand. The substance of your request "
        "is typed on Attachment 2a (documents requested), Attachment 3 "
        "(good cause), and Attachment 4 (materiality), appended behind the "
        "form."
    )
    _default_requests = _DEFAULT_SUBPOENA_REQUESTS

    _sub_def = manual_defendants[0] if manual_defendants else {}
    _sub_def_name = (_sub_def.get("name") or "").strip()

    with st.container(border=True):
        st.markdown("**Case caption**")
        st.caption(
            "Auto-filled from the plaintiff and defendant sections above — "
            "edit either field if the attachment caption should read differently."
        )
        _cap_kp = f"subcap_{name.strip()}_{_sub_def_name}"
        cap1, cap2 = st.columns(2)
        with cap1:
            sub_plaintiff_name = st.text_input(
                "Plaintiff (top of attachment)",
                value=name.strip(),
                key=f"{_cap_kp}_plaintiff",
            )
        with cap2:
            sub_defendant_name = st.text_input(
                "Defendant (top of attachment)",
                value=_sub_def_name,
                key=f"{_cap_kp}_defendant",
            )

        st.markdown("**Who are you subpoenaing?**")
        st.caption(
            "Enter the contact information of the person or agency you want "
            "records from. This is separate from the defendant — it can be a "
            "police department, city agency, business, or individual."
        )
        sub_to = st.text_input(
            "Name of person or agency to subpoena",
            key="sub_recipient_name",
            placeholder="Oakland Police Department Records Division",
        )
        sub_custodian = st.text_input(
            "Custodian of records (person or division responsible for the records)",
            key="sub_recipient_custodian",
            placeholder="Records Division",
        )
        sub_service = st.text_input(
            "Service address (where the subpoena will be delivered)",
            key="sub_recipient_service",
            placeholder="1515 Clay St, Oakland CA 94612",
        )

        st.markdown("**Documents requested**")
        st.caption(
            "What documents would you like subpoenaed for your case? "
            "Check any that apply, and add your own below. Leave blank to skip."
        )
        subpoena_checks = {}
        for _i_req, req in enumerate(_default_requests):
            subpoena_checks[req] = st.checkbox(req, value=True, key=f"sub_req_{_i_req}")
        sub_extra = st.text_area(
            "Any additional documents to request (one per line)",
            key="sub_extra_requests",
            placeholder="Any other records or documents…",
            height=80,
        )

        st.markdown("**Why the court should order production**")
        st.caption(
            "These go on Attachment 3 and Attachment 4. Defaults are written "
            "for encampment sweep cases — edit them to fit your case."
        )
        sub_good_cause = st.text_area(
            "Good cause for producing these records (Attachment 3)",
            value=_SC107_DEFAULT_GOOD_CAUSE,
            key="sub_good_cause",
            height=110,
        )
        sub_materiality = st.text_area(
            "Why these records are material to your case (Attachment 4)",
            value=_SC107_DEFAULT_MATERIALITY,
            key="sub_materiality",
            height=110,
        )
        sub_case_number = st.text_input(
            "Case number (shown on the attachments; leave blank if not yet assigned)",
            key="sub_case_number",
        )

        if st.button(
            "Generate Subpoena Package (SC-107 + Attachments)",
            use_container_width=True,
            key="gen_sc107_package",
        ):
            _sub_case = {
                "plaintiff": {"name": sub_plaintiff_name.strip() or "Plaintiff"},
                "defendant": {**_sub_def, "name": sub_defendant_name.strip() or _sub_def_name or "Defendant"},
                "case_number": sub_case_number.strip(),
                "court": manual_court,
                "subpoena": {
                    "to":               sub_to.strip(),
                    "custodian":        sub_custodian.strip(),
                    "service_location": sub_service.strip(),
                    "requests": (
                        [r for r, checked in subpoena_checks.items() if checked]
                        + [line.strip() for line in sub_extra.splitlines() if line.strip()]
                    )[:10],
                    "good_cause":  sub_good_cause.strip(),
                    "materiality": sub_materiality.strip(),
                },
            }
            try:
                with tempfile.TemporaryDirectory() as _td, _quiet():
                    _sub_out = Path(_td) / "sc107.pdf"
                    fill_sc107(_sub_case, str(_TPL / "sc107.pdf"), str(_sub_out))
                    st.session_state["sc107_package_bytes"] = _sub_out.read_bytes()
                st.session_state["sc107_package_name"] = (
                    f"{_slug(name.strip() or 'subpoena')}_sc107.pdf"
                )
            except Exception as e:
                st.session_state.pop("sc107_package_bytes", None)
                st.error(f"Could not generate the SC-107 package: {e}")

        if st.session_state.get("sc107_package_bytes"):
            st.download_button(
                "⬇️ Download SC-107 + Attachments (PDF)",
                data=st.session_state["sc107_package_bytes"],
                file_name=st.session_state.get("sc107_package_name", "sc107.pdf"),
                mime="application/pdf",
                use_container_width=True,
                key="dl_sc107_package",
            )

    # ── Helper authorization (SC-109) — separate, standalone form ───────
    st.divider()
    st.subheader("Helper Authorization (SC-109)")
    st.caption(
        "If you cannot properly present your claim on your own, a friend, "
        "family member, outreach worker, or advocate can ask the court for "
        "permission to assist you at trial (Code Civ. Proc. § 116.540). "
        "This fills the helper's information and the request to assist "
        "(item 4 of the form); the helper files it with the small claims "
        "clerk at or before the trial and signs it there."
    )
    with st.container(border=True):
        sc109_c1, sc109_c2 = st.columns(2)
        with sc109_c1:
            sc109_helper_name = st.text_input(
                "Helper's name", key="sc109_helper_name",
            )
        with sc109_c2:
            sc109_helper_rel = st.text_input(
                "Helper's relationship to you",
                key="sc109_helper_rel",
                placeholder="Friend / outreach worker / advocate",
            )
        sc109_helper_addr = st.text_input(
            "Helper's address", key="sc109_helper_addr",
        )
        sc109_reason = st.text_area(
            "Why do you need assistance presenting your case? "
            "(This goes on the form, which is not confidential.)",
            key="sc109_reason",
            placeholder=(
                "e.g., I have a disability that makes it difficult for me to "
                "speak in court and keep track of documents…"
            ),
            height=100,
        )
        sc109_c3, sc109_c4 = st.columns(2)
        with sc109_c3:
            sc109_case_number = st.text_input(
                "Case number (if assigned)", key="sc109_case_number",
            )
        with sc109_c4:
            sc109_date = st.text_input(
                "Date the helper signs (MM/DD/YYYY)", key="sc109_date",
            )

        if st.button(
            "Generate Helper Authorization (SC-109)",
            use_container_width=True,
            key="gen_sc109",
        ):
            _sc109_case = {
                "plaintiff": {"name": name.strip() or "Plaintiff"},
                "defendant": {**_sub_def, "name": _sub_def_name or "Defendant"},
                "case_number": sc109_case_number.strip(),
                "court": manual_court,
                "assistant": {
                    "name":         sc109_helper_name.strip(),
                    "address":      sc109_helper_addr.strip(),
                    "relationship": sc109_helper_rel.strip(),
                    "reason":       sc109_reason.strip(),
                    "date":         sc109_date.strip(),
                },
            }
            try:
                with tempfile.TemporaryDirectory() as _td, _quiet():
                    _sc109_out = Path(_td) / "sc109.pdf"
                    fill_sc109(_sc109_case, str(_TPL / "sc109.pdf"), str(_sc109_out))
                    st.session_state["sc109_bytes"] = _sc109_out.read_bytes()
                st.session_state["sc109_name"] = (
                    f"{_slug(name.strip() or 'helper')}_sc109.pdf"
                )
            except Exception as e:
                st.session_state.pop("sc109_bytes", None)
                st.error(f"Could not generate the SC-109: {e}")

        if st.session_state.get("sc109_bytes"):
            st.download_button(
                "⬇️ Download SC-109 Authorization to Appear",
                data=st.session_state["sc109_bytes"],
                file_name=st.session_state.get("sc109_name", "sc109.pdf"),
                mime="application/pdf",
                use_container_width=True,
                key="dl_sc109",
            )

    # ── Postpone trial (SC-150) — standalone form ────────────────────────
    st.divider()
    st.subheader("Postpone Trial (SC-150)")
    st.caption(
        "If you cannot attend your scheduled trial, ask the court to "
        "postpone it (Code Civ. Proc. § 116.570). File this at least "
        "**10 days** before trial if possible — there is a $10 fee unless "
        "the court has granted a fee waiver. Your information is filled "
        "in automatically from the plaintiff section above."
    )
    with st.container(border=True):
        sc150_role = st.radio(
            "I am the …",
            ["Plaintiff", "Defendant"],
            horizontal=True,
            key="sc150_role",
        )
        sc150_c1, sc150_c2 = st.columns(2)
        with sc150_c1:
            sc150_trial_date = st.text_input(
                "Current trial date *", key="sc150_trial_date",
                placeholder="MM/DD/YYYY",
            )
        with sc150_c2:
            sc150_new_date = st.text_input(
                "Postpone trial until (approximate date) *",
                key="sc150_new_date", placeholder="MM/DD/YYYY",
            )
        sc150_reason = st.text_area(
            "Why do you need the postponement? *",
            key="sc150_reason",
            placeholder=(
                "e.g., I am scheduled for a medical procedure that week / "
                "I am still waiting for subpoenaed records the City has not "
                "yet produced / my witness is unavailable…"
            ),
            height=100,
        )
        sc150_late_reason = st.text_area(
            "If your trial is within the next 10 days — why didn't you ask sooner?",
            key="sc150_late_reason",
            placeholder="Leave blank if your trial is more than 10 days away.",
            height=70,
        )

        st.markdown("**Has your claim been served?** (item 6 on the form)")
        _SC150_SERVICE_OPTIONS = {
            "Yes — the other parties have been served": "served",
            "No — I am a defendant and have not filed a claim": "not_filed",
            "No — some parties have not been served": "not_served",
            "I don't know — the court clerk mailed my claim": "unknown",
        }
        sc150_service_label = st.selectbox(
            "Service status",
            list(_SC150_SERVICE_OPTIONS),
            key="sc150_service_status",
        )
        sc150_status = _SC150_SERVICE_OPTIONS[sc150_service_label]

        sc150_served, sc150_unserved, sc150_unknown = [], [], []
        if sc150_status == "served":
            _sc150_def_name = (_sub_def.get("name") or "").strip()
            for _si in (1, 2):
                sv1, sv2, sv3 = st.columns([2, 1, 1])
                with sv1:
                    _sname = st.text_input(
                        f"Served party {_si} — name",
                        value=_sc150_def_name if _si == 1 else "",
                        key=f"sc150_served_name{_si}",
                    )
                with sv2:
                    _scounty = st.text_input(
                        "County they live in",
                        value=manual_court.get("county", "") if _si == 1 else "",
                        key=f"sc150_served_county{_si}",
                    )
                with sv3:
                    _sdate = st.text_input(
                        "Date served", placeholder="MM/DD/YYYY",
                        key=f"sc150_served_date{_si}",
                    )
                if _sname.strip():
                    sc150_served.append({
                        "name":   _sname.strip(),
                        "county": _scounty.strip(),
                        "date":   _sdate.strip(),
                    })
        elif sc150_status == "not_served":
            sc150_unserved = [
                n.strip() for n in st.text_input(
                    "Parties not yet served (separate names with a semicolon)",
                    key="sc150_unserved_names",
                ).split(";") if n.strip()
            ]
        elif sc150_status == "unknown":
            sc150_unknown = [
                n.strip() for n in st.text_input(
                    "Parties whose service receipt is unconfirmed (separate with a semicolon)",
                    key="sc150_unknown_names",
                ).split(";") if n.strip()
            ]

        sc150_c3, sc150_c4 = st.columns(2)
        with sc150_c3:
            sc150_case_number = st.text_input(
                "Case number", key="sc150_case_number",
            )
        with sc150_c4:
            sc150_sign_date = st.text_input(
                "Date you sign the request (MM/DD/YYYY)", key="sc150_sign_date",
            )

        if st.button(
            "Generate Request to Postpone Trial (SC-150)",
            use_container_width=True,
            key="gen_sc150",
        ):
            _sc150_case = {
                "plaintiff": {
                    "name":   name.strip() or "Plaintiff",
                    "street": street.strip(),
                    "city":   city.strip(),
                    "state":  state.strip() or "CA",
                    "zip":    zip_.strip(),
                    "phone":  phone.strip(),
                },
                "defendant": {**_sub_def, "name": _sub_def.get("name") or "Defendant"},
                "case_number": sc150_case_number.strip(),
                "court": manual_court,
                "postponement": {
                    "requester_name":     name.strip(),
                    "role":               sc150_role.lower(),
                    "phone":              phone.strip(),
                    "current_trial_date": sc150_trial_date.strip(),
                    "requested_date":     sc150_new_date.strip(),
                    "reason":             sc150_reason.strip(),
                    "late_reason":        sc150_late_reason.strip(),
                    "service_status":     sc150_status,
                    "served":             sc150_served,
                    "unserved_names":     sc150_unserved,
                    "unknown_names":      sc150_unknown,
                    "request_date":       sc150_sign_date.strip(),
                },
            }
            if not sc150_reason.strip() or not sc150_new_date.strip():
                st.error(
                    "Please fill in the requested new trial date and the "
                    "reason for the postponement."
                )
            else:
                try:
                    with tempfile.TemporaryDirectory() as _td, _quiet():
                        _sc150_out = Path(_td) / "sc150.pdf"
                        fill_sc150(_sc150_case, str(_TPL / "sc150.pdf"), str(_sc150_out))
                        st.session_state["sc150_bytes"] = _sc150_out.read_bytes()
                    st.session_state["sc150_name"] = (
                        f"{_slug(name.strip() or 'postpone')}_sc150.pdf"
                    )
                except Exception as e:
                    st.session_state.pop("sc150_bytes", None)
                    st.error(f"Could not generate the SC-150: {e}")

        if st.session_state.get("sc150_bytes"):
            st.download_button(
                "⬇️ Download SC-150 Request to Postpone Trial",
                data=st.session_state["sc150_bytes"],
                file_name=st.session_state.get("sc150_name", "sc150.pdf"),
                mime="application/pdf",
                use_container_width=True,
                key="dl_sc150",
            )
            st.caption(
                "File the signed form with the small claims clerk (with the "
                "$10 fee unless waived) and **mail or deliver a copy to every "
                "other party in the case**. The court will notify you whether "
                "the postponement is granted."
            )



# ══════════════════════════════════════════════════════
# TAB 2 — SPREADSHEET IMPORT
# ══════════════════════════════════════════════════════

with tab_sheet:
    st.subheader("Batch Import from Spreadsheet")

    # ── Template downloads ─────────────────────────────────────────────────
    c_info, c_tmpl = st.columns([3, 1])
    with c_info:
        st.info(
            "**Oakland intake format** (columns: Name, Address, Phone Number, "
            "Location of Injury, Date of Injury) is auto-detected from your Google Sheet. "
            "You'll set claim amount, filing dates, and fee waiver defaults that apply to "
            "every client in the batch."
        )
    with c_tmpl:
        st.download_button(
            "📥 Download Full Template CSV",
            data=_csv_template_bytes(),
            file_name="cases_template.csv",
            mime="text/csv",
            width="stretch",
            help="Use this template if you want to specify all fields per client.",
        )

    # ── Column reference ───────────────────────────────────────────────────
    with st.expander("View full template column reference"):
        col_df = pd.DataFrame(
            [(c[0], "✓" if c[2] else "", c[1], c[3]) for c in _TEMPLATE_COLS],
            columns=["Column", "Required", "Description", "Example"],
        )
        st.dataframe(col_df, use_container_width=True, hide_index=True)

    with st.expander("SC-107 subpoena checklist and spreadsheet fields"):
        subpoena_df = pd.DataFrame([
            ["subpoena_case_caption",   "Case caption for SC-107 subpoena"],
            ["subpoena_to",             "Subpoena recipient or agency to which records are directed"],
            ["subpoena_custodian",      "Custodian of records responsible for producing documents"],
            ["subpoena_service_location","Service address where the subpoena may be delivered"],
            ["subpoena_request_1",      "Body-worn camera and officer dashboard footage from the sweep"],
            ["subpoena_request_2",      "Incident reports, notes, and supplemental reports related to the sweep"],
            ["subpoena_request_3",      "Dispatch logs, radio transmissions, and 911/311 recordings"],
            ["subpoena_request_4",      "Officer complaint, investigation, and disciplinary records"],
            ["subpoena_request_5",      "Internal communications, emails, memos, and directives about encampment sweeps"],
            ["subpoena_request_6",      "Policies, training materials, use-of-force guidelines, and encampment protocols"],
            ["subpoena_request_7",      "Property seizure, storage, chain-of-custody, and disposal records"],
            ["subpoena_request_8",      "Surveillance camera and private video footage from the sweep location"],
            ["subpoena_request_9",      "Records of coordination between police, DPW, and other agencies"],
            ["subpoena_request_10",     "Logs, schedules, and written directives authorizing the sweeps"],
        ], columns=["Spreadsheet Column", "Request description"])
        st.dataframe(subpoena_df, use_container_width=True, hide_index=True)

    # ── Upload ─────────────────────────────────────────────────────────────
    uploaded = st.file_uploader(
        "Upload spreadsheet (CSV or XLSX)",
        type=["csv", "xlsx"],
        label_visibility="collapsed",
    )

    if not uploaded:
        st.stop()

    try:
        if uploaded.name.endswith(".xlsx"):
            df = pd.read_excel(uploaded, dtype=str)
        else:
            df = pd.read_csv(uploaded, dtype=str)
        df = df.fillna("")
        # Drop rows where Name/plaintiff_name is blank
        name_col = "Name" if "Name" in df.columns else "plaintiff_name"
        df = df[df[name_col].str.strip() != ""].reset_index(drop=True)
    except Exception as e:
        st.error(f"Could not read file: {e}")
        st.stop()

    fmt = _detect_format(df)
    st.write(f"**{len(df)} client(s) found** — format: `{fmt}`")

    # ── Preview ────────────────────────────────────────────────────────────
    preview_cols = {
        "oakland_intake": ["Name", "Address", "Phone Number", "Location of Injury", "Date of Injury"],
        "template": [c[0] for c in _TEMPLATE_COLS if c[2]],  # required cols only
    }
    show_cols = [c for c in preview_cols.get(fmt, list(df.columns)) if c in df.columns]
    st.dataframe(df[show_cols] if show_cols else df, use_container_width=True, height=250)

    if fmt == "unknown":
        st.warning(
            "Column names not recognized. Rename columns to match the Oakland intake format "
            "(Name, Address, Phone Number, Location of Injury, Date of Injury) "
            "or the full template format (plaintiff_name, etc.)."
        )
        st.stop()

    # ═══════════════════════════════════════════════════════
    # OAKLAND INTAKE FORMAT: batch defaults form
    # ═══════════════════════════════════════════════════════
    if fmt == "oakland_intake":
        st.divider()
        st.subheader("Batch Settings")
        st.caption(
            "These values apply to **all clients** in the spreadsheet. "
            "They fill in the fields not captured in the Oakland intake sheet."
        )

        st.markdown("**Filing Court**")
        batch_court = _court_selector("batch")
        st.caption(
            f"Court: **Superior Court of California, County of {batch_court['county']}** · "
            f"{batch_court['address']}, {batch_court['city']}, CA {batch_court['zip']}"
        )

        st.markdown("**Defendant** — applies to every row in the spreadsheet")
        batch_def = _defendant_block("batch", 0, is_primary=True)

        with st.container(border=True):
            d1, d2 = st.columns(2)
            with d1:
                b_filing_date      = st.text_input("Filing Date *", placeholder="MM/DD/YYYY",
                                                    help="Date you're filing the small claims paperwork.")
                b_govt_claim_date  = st.text_input("Govt Claim Filed with City Clerk *",
                                                    placeholder="MM/DD/YYYY",
                                                    help="Date the government tort claim was filed.")
                b_claim_amount     = st.text_input("Claim Amount ($) per client", value="10000",
                                                    help="Max $12,500 for individuals.")
            with d2:
                b_fw_basis = st.radio(
                    "Fee Waiver Basis",
                    ["5c — Cannot afford fees", "5a — Public benefits", "5b — Income threshold"],
                    horizontal=True,
                )
                b_recv_medi_cal = st.checkbox("All clients receive Medi-Cal")
                b_recv_snap     = st.checkbox("All clients receive CalFresh / SNAP")

            st.markdown("**Income & Expenses (applies to all — edit per-client if needed)**")
            e1, e2, e3, e4 = st.columns(4)
            with e1:
                b_income_source = st.text_input("Income Source", placeholder="General Assistance")
                b_income_amount = st.text_input("Monthly Income ($)", placeholder="400")
                b_total_income  = st.text_input("Total Monthly Income ($)", placeholder="400")
            with e2:
                b_exp_food    = st.text_input("Food ($)", value="0")
                b_exp_medical = st.text_input("Medical ($)", value="0")
            with e3:
                b_exp_transport = st.text_input("Transport ($)", value="0")
                b_exp_housing   = st.text_input("Housing ($)", value="0")
            with e4:
                b_total_expenses = st.text_input("Total Expenses ($)", placeholder="300")

            st.markdown("**Claim Narrative** (optional — leave blank for auto-generated text)")
            b_claim_reason = st.text_area(
                "Claim reason template",
                placeholder=(
                    "Leave blank to auto-generate: 'On [date], City of Oakland DPW "
                    "swept [location]…'  — incident date and location are filled "
                    "from each client's row."
                ),
                height=80,
            )
            b_declaration = st.text_area(
                "Declaration template",
                placeholder="Leave blank to auto-generate a first-person declaration.",
                height=80,
            )

            run_batch = st.button(
                "Generate All Forms", type="primary", use_container_width=True
            )

        if run_batch:
            _batch_defendant = batch_def if (batch_def or {}).get("name") else DEFENDANT_DEFAULTS["city_of_oakland"]
            defaults = {
                "court":                 batch_court,
                "defendant":             _batch_defendant,
                "filing_date":           b_filing_date.strip(),
                "govt_claim_filed_date": b_govt_claim_date.strip(),
                "claim_amount":          b_claim_amount.strip(),
                "fw_basis":              b_fw_basis.split(" — ")[0].strip(),
                "receives_medi_cal":     b_recv_medi_cal,
                "receives_snap":         b_recv_snap,
                "receives_calworks":     False,
                "income_source":         b_income_source.strip(),
                "income_amount":         b_income_amount.strip(),
                "total_income":          b_total_income.strip(),
                "expense_food":          b_exp_food.strip(),
                "expense_medical":       b_exp_medical.strip(),
                "expense_transport":     b_exp_transport.strip(),
                "expense_housing":       b_exp_housing.strip(),
                "total_expenses":        b_total_expenses.strip(),
                "claim_reason":          b_claim_reason.strip(),
                "declaration":           b_declaration.strip(),
                "damages_calculation":   "",
            }

            results = []
            progress = st.progress(0, text="Generating forms…")
            for i, (_, row) in enumerate(df.iterrows()):
                pname = str(row.get("Name", f"Row {i+1}")).strip()
                try:
                    case = intake_row_to_case(row, defaults)
                    try:
                        _capture_case_record(case)
                    except Exception:
                        pass  # capture failure shouldn't block form generation
                    pdfs = _generate_pdfs(case)
                    results.append((pname, pdfs, None))
                except ValueError as e:
                    results.append((pname, None, str(e)))
                except Exception as e:
                    results.append((pname, None, f"Unexpected error: {e}"))
                progress.progress((i + 1) / len(df), text=f"Processed {i+1}/{len(df)}…")
            progress.empty()

            ok   = [(n, p, _) for n, p, _ in results if _ is None]
            fail = [(n, p, e) for n, p, e in results if e is not None]

            if ok:
                st.success(f"Generated forms for **{len(ok)}** of {len(results)} clients.")
                zip_buf = io.BytesIO()
                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for pname, pdfs, _ in ok:
                        slug = _slug(pname)
                        for lbl, data in pdfs.items():
                            zf.writestr(
                                f"{slug}/{slug}_{lbl.lower().replace('-','')}.pdf", data
                            )
                st.download_button(
                    "⬇️  Download All Clients (ZIP)",
                    data=zip_buf.getvalue(),
                    file_name="oakland_encampment_forms.zip",
                    mime="application/zip",
                    type="primary",
                    width="stretch",
                )

            if fail:
                st.warning(f"{len(fail)} client(s) could not be processed:")
                for pname, _, err in fail:
                    with st.expander(f"Error — {pname}"):
                        st.text(err)

    # ═══════════════════════════════════════════════════════
    # TEMPLATE FORMAT: direct processing
    # ═══════════════════════════════════════════════════════
    elif fmt == "template":
        if st.button("Generate All Forms", type="primary", width="stretch"):
            results = []
            progress = st.progress(0, text="Generating forms…")
            for i, (_, row) in enumerate(df.iterrows()):
                pname = str(row.get("plaintiff_name", f"Row {i+1}")).strip()
                try:
                    case = template_row_to_case(row)
                    try:
                        _capture_case_record(case)
                    except Exception:
                        pass  # capture failure shouldn't block form generation
                    pdfs = _generate_pdfs(case)
                    results.append((pname, pdfs, None))
                except ValueError as e:
                    results.append((pname, None, str(e)))
                except Exception as e:
                    results.append((pname, None, f"Unexpected error: {e}"))
                progress.progress((i + 1) / len(df), text=f"Processed {i+1}/{len(df)}…")
            progress.empty()

            ok   = [(n, p, _) for n, p, _ in results if _ is None]
            fail = [(n, p, e) for n, p, e in results if e is not None]

            if ok:
                st.success(f"Generated forms for {len(ok)} of {len(results)} clients.")
                zip_buf = io.BytesIO()
                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for pname, pdfs, _ in ok:
                        slug = _slug(pname)
                        for lbl, data in pdfs.items():
                            zf.writestr(
                                f"{slug}/{slug}_{lbl.lower().replace('-','')}.pdf", data
                            )
                st.download_button(
                    "⬇️  Download All Clients (ZIP)",
                    data=zip_buf.getvalue(),
                    file_name="all_cases.zip",
                    mime="application/zip",
                    type="primary",
                    width="stretch",
                )

            if fail:
                st.warning(f"{len(fail)} client(s) had errors:")
                for pname, _, err in fail:
                    with st.expander(f"Error — {pname}"):
                        st.text(err)


